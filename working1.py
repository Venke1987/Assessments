import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import random
import openai
import json
from datetime import datetime
import fitz  # PyMuPDF for PDFs
import docx  # For Word document handling
import ast  # For checking Python syntax
import io  # For handling uploaded files
import sqlite3
import os
from io import BytesIO
from dotenv import load_dotenv
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity




# Load environment variables (if needed)
load_dotenv()

# openai.api_key = os.getenv("OPENAI_API_KEY")
openai.api_key = st.secrets["OPENAI_API_KEY"]
client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])  # Use Streamlit Secrets API Key


##############################################################################
#                         Database Initialization                            #
##############################################################################
def init_db():
    """
    Initialize (or connect to) the SQLite database for storing quiz and plagiarism checks.
    """
    conn = sqlite3.connect('student_quiz_history.db')
    c = conn.cursor()
    c.execute('''
    CREATE TABLE IF NOT EXISTS quiz_results (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        student_id TEXT NOT NULL,
        student_name TEXT NOT NULL,
        topic TEXT NOT NULL,
        score INTEGER NOT NULL,
        total_questions INTEGER NOT NULL,
        timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    conn.commit()
    return conn

def save_quiz_result(student_id, student_name, topic, score, total_questions):
    """
    Save the quiz result into the quiz_results table.
    """
    conn = init_db()
    c = conn.cursor()
    c.execute('''
    INSERT INTO quiz_results (student_id, student_name, topic, score, total_questions, timestamp)
    VALUES (?, ?, ?, ?, ?, ?)
    ''', (student_id, student_name, topic, score, total_questions, datetime.now()))
    conn.commit()
    conn.close()

def get_student_quiz_history(student_id=None):
    """
    Retrieve quiz history for a student (or all students if no ID passed).
    Return a pandas DataFrame with results.
    """
    conn = init_db()
    c = conn.cursor()
    
    if student_id:
        c.execute('''
        SELECT student_id, student_name, topic, score, total_questions, timestamp
        FROM quiz_results
        WHERE student_id = ?
        ORDER BY timestamp DESC
        ''', (student_id,))
    else:
        c.execute('''
        SELECT student_id, student_name, topic, score, total_questions, timestamp
        FROM quiz_results
        ORDER BY timestamp DESC
        ''')
    
    results = c.fetchall()
    conn.close()
    
    if results:
        return pd.DataFrame(results, columns=['student_id', 'student_name', 'topic', 'score', 'total_questions', 'timestamp'])
    return pd.DataFrame()

##############################################################################
#                           File Extraction Helpers                          #
##############################################################################
def extract_text_from_docx(uploaded_file):
    """Extract text from an in-memory docx file."""
    try:
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def extract_text_from_pdf(uploaded_file):
    """Extract text from an in-memory pdf file."""
    try:
        text = ""
        with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text() + "\n"
        return text
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def check_python_syntax(code):
    """Check for syntax errors in uploaded Python code."""
    try:
        ast.parse(code)
        return None  # No syntax errors
    except SyntaxError as e:
        return f"Syntax Error: {e}"

##############################################################################
#             Local Plagiarism Check: Compare with Past Reports             #
##############################################################################
def extract_text_from_local_file(file_path):
    """
    Extract text from a local PDF or DOCX file using PyMuPDF (fitz) or python-docx.
    """
    text = ""
    try:
        if file_path.lower().endswith(".pdf"):
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text += page.get_text()
        elif file_path.lower().endswith(".docx"):
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        text = f"Error extracting text from {file_path}: {str(e)}"
    return text

def compute_local_plagiarism_scores(student_text, folder_path="local_reports"):
    """
    Compare `student_text` with each file in `folder_path` and return {filename: similarity_score}.
    Uses TF-IDF + Cosine Similarity.
    """
    similarities = {}
    if not os.path.exists(folder_path):
        return {"Error": f"Folder '{folder_path}' not found. Please create it and add your reports."}
    for fname in os.listdir(folder_path):
        if fname.lower().endswith((".pdf", ".docx")):
            file_path = os.path.join(folder_path, fname)
            local_text = extract_text_from_local_file(file_path)
            if local_text.startswith("Error extracting text"):
                # If extraction failed, just record the error
                similarities[fname] = local_text
            else:
                # TF-IDF Cosine similarity
                docs = [student_text, local_text]
                tfidf = TfidfVectorizer().fit_transform(docs)
                score = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
                similarities[fname] = score
    return similarities

##############################################################################
#            Sample Student Data & Questions (As Provided Originally)        #
##############################################################################
students_data = {
    "SEEE001": {"name": "Adhithya V", "proficiency": {"Fuzzy Logic": {1: 0.59, 2: 0.64, 3: 0.69, 4: 0.74, 5: 0.79, 6: 0.84}}},
    "SEEE002": {"name": "Akety Manjunath", "proficiency": {"Data Science": {1: 0.6, 2: 0.65, 3: 0.7, 4: 0.75, 5: 0.8, 6: 0.85}}},
    "SEEE003": {"name": "Aravind S", "proficiency": {"Bayesian Learning": {1: 0.37, 2: 0.42, 3: 0.47, 4: 0.52, 5: 0.57, 6: 0.62}}},
    "SEEE004": {"name": "Aswin S", "proficiency": {"NLP": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}},
    "SEEE005": {"name": "Avinash Kumar R", "proficiency": {"Data Science": {1: 0.69, 2: 0.74, 3: 0.79, 4: 0.84, 5: 0.89, 6: 0.94}}},
    "SEEE007": {"name": "Bhavya Sri B", "proficiency": {"IoT": {1: 0.59, 2: 0.64, 3: 0.69, 4: 0.74, 5: 0.79, 6: 0.84}}},
    "SEEE008": {"name": "Challagandla Anantha Pavan", "proficiency": {"Cybersecurity": {1: 0.7, 2: 0.75, 3: 0.8, 4: 0.85, 5: 0.9, 6: 0.95}}},
    "SEEE009": {"name": "Nagadarahas Kumar C S", "proficiency": {"Cybersecurity": {1: 0.6, 2: 0.65, 3: 0.7, 4: 0.75, 5: 0.8, 6: 0.85}}},
    "SEEE010": {"name": "Dodda Sri Pujitha", "proficiency": {"Edge AI": {1: 0.73, 2: 0.78, 3: 0.83, 4: 0.88, 5: 0.93, 6: 0.98}}},
    "SEEE011": {"name": "Dondapati Vallapala Yami", "proficiency": {"Deep Learning": {1: 0.58, 2: 0.63, 3: 0.68, 4: 0.73, 5: 0.78, 6: 0.83}}},
    "SEEE012": {"name": "Guduru Venkata Sai Karthikeya", "proficiency": {"Generative AI": {1: 0.42, 2: 0.47, 3: 0.52, 4: 0.57, 5: 0.62, 6: 0.67}}},
    "SEEE013": {"name": "Hamsitha P", "proficiency": {"AI Ethics": {1: 0.39, 2: 0.44, 3: 0.49, 4: 0.54, 5: 0.59, 6: 0.64}}},
    "SEEE014": {"name": "Harini M", "proficiency": {"Generative AI": {1: 0.5, 2: 0.55, 3: 0.6, 4: 0.65, 5: 0.7, 6: 0.75}}},
    "SEEE015": {"name": "Harrish B", "proficiency": {"Robotics": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE016": {"name": "Jivan Prasath S", "proficiency": {"Swarm Intelligence": {1: 0.51, 2: 0.56, 3: 0.61, 4: 0.66, 5: 0.71, 6: 0.76}}},
    "SEEE017": {"name": "Jonnalagadda Susrith", "proficiency": {"Quantum Computing": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE018": {"name": "Kamaleshwar M", "proficiency": {"Data Science": {1: 0.72, 2: 0.77, 3: 0.82, 4: 0.87, 5: 0.92, 6: 0.97}}},
    "SEEE019": {"name": "Karthik Periyakarupphan M", "proficiency": {"NLP": {1: 0.36, 2: 0.41, 3: 0.46, 4: 0.51, 5: 0.56, 6: 0.61}}},
    "SEEE020": {"name": "Kathirazagan V", "proficiency": {"Bayesian Learning": {1: 0.44, 2: 0.49, 3: 0.54, 4: 0.59, 5: 0.64, 6: 0.69}}},
    "SEEE021": {"name": "Kishore R", "proficiency": {"Robotics": {1: 0.72, 2: 0.77, 3: 0.82, 4: 0.87, 5: 0.92, 6: 0.97}}},
    "SEEE022": {"name": "Krishnakumar Aditi J", "proficiency": {"Computer Vision": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE023": {"name": "Logavarshini K", "proficiency": {"Optimization": {1: 0.7, 2: 0.75, 3: 0.8, 4: 0.85, 5: 0.9, 6: 0.95}}},
    "SEEE024": {"name": "Malapalli Charitha Reddy", "proficiency": {"Deep Learning": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}},
    "SEEE025": {"name": "Manas M Nair", "proficiency": {"Bayesian Learning": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE026": {"name": "Mervath M", "proficiency": {"Robotics": {1: 0.55, 2: 0.6, 3: 0.65, 4: 0.7, 5: 0.75, 6: 0.8}}},
    "SEEE027": {"name": "Nidhish Kumar K", "proficiency": {"Bayesian Learning": {1: 0.37, 2: 0.42, 3: 0.47, 4: 0.52, 5: 0.57, 6: 0.62}}},
    "SEEE028": {"name": "Palapati Jahnavi", "proficiency": {"Reinforcement Learning": {1: 0.65, 2: 0.7, 3: 0.75, 4: 0.8, 5: 0.85, 6: 0.9}}},
    "SEEE029": {"name": "Pola Kaarthi", "proficiency": {"Fuzzy Logic": {1: 0.74, 2: 0.79, 3: 0.84, 4: 0.89, 5: 0.94, 6: 0.99}}},
    "SEEE030": {"name": "Pullela Vaishnavi", "proficiency": {"Data Science": {1: 0.39, 2: 0.44, 3: 0.49, 4: 0.54, 5: 0.59, 6: 0.64}}},
    "SEEE031": {"name": "Raghul T", "proficiency": {"Generative AI": {1: 0.43, 2: 0.48, 3: 0.53, 4: 0.58, 5: 0.63, 6: 0.68}}},
    "SEEE032": {"name": "Rajaprabu K", "proficiency": {"Big Data": {1: 0.41, 2: 0.46, 3: 0.51, 4: 0.56, 5: 0.61, 6: 0.66}}},
    "SEEE033": {"name": "Rishi A", "proficiency": {"NLP": {1: 0.41, 2: 0.46, 3: 0.51, 4: 0.56, 5: 0.61, 6: 0.66}}},
    "SEEE034": {"name": "Rithick Reddy S", "proficiency": {"Swarm Intelligence": {1: 0.48, 2: 0.53, 3: 0.58, 4: 0.63, 5: 0.68, 6: 0.73}}},
    "SEEE035": {"name": "Tera Sai Tejeshwar Reddy", "proficiency": {"Swarm Intelligence": {1: 0.59, 2: 0.64, 3: 0.69, 4: 0.74, 5: 0.79, 6: 0.84}}},
    "SEEE036": {"name": "Saikirthiga R", "proficiency": {"Data Science": {1: 0.55, 2: 0.6, 3: 0.65, 4: 0.7, 5: 0.75, 6: 0.8}}},
    "SEEE037": {"name": "Sangam JayaVardhan Reddy", "proficiency": {"Big Data": {1: 0.38, 2: 0.43, 3: 0.48, 4: 0.53, 5: 0.58, 6: 0.63}}},
    "SEEE038": {"name": "Sarveshwaran S", "proficiency": {"ML Basics": {1: 0.71, 2: 0.76, 3: 0.81, 4: 0.86, 5: 0.91, 6: 0.96}}},
    "SEEE039": {"name": "Shanjay Sundhar S", "proficiency": {"IoT": {1: 0.68, 2: 0.73, 3: 0.78, 4: 0.83, 5: 0.88, 6: 0.93}}},
    "SEEE040": {"name": "Shreeya Kannan", "proficiency": {"Fuzzy Logic": {1: 0.62, 2: 0.67, 3: 0.72, 4: 0.77, 5: 0.82, 6: 0.87}}},
    "SEEE041": {"name": "Shreya S", "proficiency": {"AI Ethics": {1: 0.61, 2: 0.66, 3: 0.71, 4: 0.76, 5: 0.81, 6: 0.86}}},
    "SEEE042": {"name": "Sri Kamal Krishank S", "proficiency": {"Data Science": {1: 0.46, 2: 0.51, 3: 0.56, 4: 0.61, 5: 0.66, 6: 0.71}}},
    "SEEE043": {"name": "Sri Ramana S", "proficiency": {"Quantum Computing": {1: 0.37, 2: 0.42, 3: 0.47, 4: 0.52, 5: 0.57, 6: 0.62}}},
    "SEEE044": {"name": "Subashree S", "proficiency": {"Reinforcement Learning": {1: 0.48, 2: 0.53, 3: 0.58, 4: 0.63, 5: 0.68, 6: 0.73}}},
    "SEEE045": {"name": "Subhiksha N", "proficiency": {"Bayesian Learning": {1: 0.38, 2: 0.43, 3: 0.48, 4: 0.53, 5: 0.58, 6: 0.63}}},
    "SEEE046": {"name": "Sukhi Sudharshan S", "proficiency": {"Robotics": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE047": {"name": "Thanuja", "proficiency": {"Computer Vision": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}},
    "SEEE048": {"name": "Vasantha Kumar A", "proficiency": {"Quantum Computing": {1: 0.5, 2: 0.55, 3: 0.6, 4: 0.65, 5: 0.7, 6: 0.75}}},
    "SEEE049": {"name": "Velmugilan S", "proficiency": {"Embedded Systems": {1: 0.52, 2: 0.57, 3: 0.62, 4: 0.67, 5: 0.72, 6: 0.77}}},
    "SEEE050": {"name": "Velmurugan S", "proficiency": {"Fuzzy Logic": {1: 0.5, 2: 0.55, 3: 0.6, 4: 0.65, 5: 0.7, 6: 0.75}}},
    "SEEE051": {"name": "Viamrsh P S", "proficiency": {"Generative AI": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE052": {"name": "Vilohitan R", "proficiency": {"Reinforcement Learning": {1: 0.44, 2: 0.49, 3: 0.54, 4: 0.59, 5: 0.64, 6: 0.69}}},
    "SEEE053": {"name": "Vundela Guru Venkata Ajay Kumar Reddy", "proficiency": {"Big Data": {1: 0.72, 2: 0.77, 3: 0.82, 4: 0.87, 5: 0.92, 6: 0.97}}},
    "SEEE054": {"name": "Yadlapalli Madhuri", "proficiency": {"Robotics": {1: 0.44, 2: 0.49, 3: 0.54, 4: 0.59, 5: 0.64, 6: 0.69}}},
    "SEEE055": {"name": "Yohan K", "proficiency": {"Data Science": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE056": {"name": "Vinupriya A", "proficiency": {"Computer Vision": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE057": {"name": "Aparajita S", "proficiency": {"Computer Vision": {1: 0.73, 2: 0.78, 3: 0.83, 4: 0.88, 5: 0.93, 6: 0.98}}},
    "SEEE058": {"name": "Sriram L", "proficiency": {"NLP": {1: 0.63, 2: 0.68, 3: 0.73, 4: 0.78, 5: 0.83, 6: 0.88}}},
    "SEEE059": {"name": "Nandhini S", "proficiency": {"Deep Learning": {1: 0.55, 2: 0.6, 3: 0.65, 4: 0.7, 5: 0.75, 6: 0.8}}},
    "SEEE060": {"name": "Sreenath G", "proficiency": {"IoT": {1: 0.52, 2: 0.57, 3: 0.62, 4: 0.67, 5: 0.72, 6: 0.77}}},
    "SEEE061": {"name": "Swaran V", "proficiency": {"AI Ethics": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}}
}

questions = [
    {"text": "What is the primary goal of supervised learning?",
     "options": [
         "To cluster data points without labels",
         "To learn patterns from labeled data to make predictions",
         "To reduce the dimensionality of data",
         "To generate new data samples"
     ],
     "correct": "To learn patterns from labeled data to make predictions"},
    {"text": "Which algorithm is used for classification?",
     "options": [
         "K-Means", "Decision Tree", "PCA", "Apriori"
     ],
     "correct": "Decision Tree"},
    {"text": "What does CNN stand for?",
     "options": [
         "Convolutional Neural Network", "Central Neural Network",
         "Computational Node Network", "Convolutional Node Network"
     ],
     "correct": "Convolutional Neural Network"},
    {"text": "Which technique helps reduce overfitting?",
     "options": [
         "Regularization", "Dropout", "Both 1 & 2", "None"
     ],
     "correct": "Both 1 & 2"}
]

##############################################################################
#                          Session State Initialization                      #
##############################################################################
if "quiz_active" not in st.session_state:
    st.session_state.quiz_active = False
if "question_count" not in st.session_state:
    st.session_state.question_count = 0
if "total_questions" not in st.session_state:
    st.session_state.total_questions = 5
if "current_question" not in st.session_state:
    st.session_state.current_question = None
if "score" not in st.session_state:
    st.session_state.score = 0

# Initialize the SQLite database
init_db()

##############################################################################
#                              Streamlit Pages                               #
##############################################################################
st.title("🚀 Generative AI-Based MEC102 Engineering Design Project Report Assessment System")
st.sidebar.header("Navigation")
page = st.sidebar.radio("Go to", ["📊 Dashboard",
                                  "🔍 Plagiarism/Reasoning Finder",
                                  "AI-Powered Storytelling",
                                  "🔍 AI Peer Assessment"])

### ---------------------- 1) DASHBOARD ---------------------- ###
if page == "📊 Dashboard":
    st.header("📊 Class Performance Dashboard")
    
    # Select Student
    student_id = st.selectbox("Select a Student", list(students_data.keys()))
    student = students_data[student_id]
    st.subheader(f"Student: {student['name']}")
    
    # Proficiency Data
    prof_data = student["proficiency"]
    topics = list(prof_data.keys())
    bloom_levels = [1, 2, 3, 4, 5, 6]
    
    # Heatmap DataFrame
    heatmap_data = pd.DataFrame(
        {topic: [prof_data[topic].get(level, 0.5) for level in bloom_levels] for topic in topics},
        index=[f"Level {lvl}" for lvl in bloom_levels]
    )
    
    st.subheader("📌 Bloom's Taxonomy Heatmap")
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.heatmap(heatmap_data, annot=True, cmap="coolwarm", linewidths=0.5, ax=ax)
    st.pyplot(fig)
    
    st.subheader("📈 Student Progress Over Time")
    dates = [datetime.now().replace(day=i) for i in range(1, 11)]
    scores = [random.uniform(0.4, 0.9) for _ in range(10)]
    plt.figure(figsize=(8, 5))
    plt.plot(dates, scores, marker='o', linestyle='-')
    plt.xticks(rotation=45)
    plt.xlabel("Date")
    plt.ylabel("Proficiency Score")
    plt.title("Student Progress")
    st.pyplot(plt)

### ---------------------- 2) PLAGIARISM/REASONING FINDER ---------------------- ###
elif page == "🔍 Plagiarism/Reasoning Finder":
    st.header("🔍 Plagiarism/Reasoning Finder")
    uploaded_file = st.file_uploader("Upload student's document (.docx or .pdf)", type=["docx", "pdf"])

    student_submission = ""
    if uploaded_file:
        file_type = uploaded_file.name.split(".")[-1].lower()
        
        if file_type == "docx":
            student_submission = extract_text_from_docx(uploaded_file)
        elif file_type == "pdf":
            student_submission = extract_text_from_pdf(uploaded_file)

        st.subheader("📄 Extracted Submission:")
        st.text_area("Extracted Text", student_submission, height=300)

    rubric = {"Concept Understanding": 10, "Implementation": 10, "Analysis": 10, "Clarity": 5, "Creativity": 5}

    col1, col2, col3 = st.columns([1, 1, 1])
    
    # -- Generate AI Assessment --
    with col1:
        if st.button("📝 Generate AI Assessment"):
            if student_submission.strip():
                prompt = (
                    f"Assess the following submission based on this rubric: {json.dumps(rubric)}\n\n"
                    f"Submission:\n{student_submission}"
                )
                
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "You are an AI grading student assignments."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    max_tokens=500
                )
                
                review = response.choices[0].message.content.strip()
                st.success("✅ AI Feedback Generated:")
                st.write(review)
            else:
                st.warning("⚠️ Please upload a valid document.")
    
    # -- LLM-Based Plagiarism Check --
    with col2:
        if st.button("🔍 Check for Plagiarism (LLM-based)"):
            if student_submission.strip():
                with st.spinner("Checking plagiarism..."):
                    plagiarism_prompt = (
                        f"Analyze the following document for plagiarism risk. "
                        f"Your response MUST include a clear plagiarism percentage in the first line with this exact format: "
                        f"'Plagiarism Risk: XX%' (where XX is a number between 0 and 100). "
                        f"Then provide your explanation below.\n\n"
                        f"Document:\n{student_submission}"
                    )

                    plagiarism_response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "You are an AI that checks for plagiarism. Always provide a clear plagiarism percentage in your first line."},
                            {"role": "user", "content": plagiarism_prompt}
                        ],
                        temperature=0.3,
                        max_tokens=400
                    )

                    plagiarism_result = plagiarism_response.choices[0].message.content.strip()
                    
                    # Try extracting plagiarism percentage
                    try:
                        if "Plagiarism Risk:" in plagiarism_result:
                            first_line = plagiarism_result.split('\n')[0]
                            percentage_text = first_line.split('Plagiarism Risk:')[1].strip()
                            if '%' in percentage_text:
                                percentage = float(percentage_text.replace('%', '').strip())
                                
                                st.subheader("Plagiarism Risk Score:")
                                st.progress(percentage/100)
                                
                                if percentage < 30:
                                    st.success(f"📊 Plagiarism Risk: {percentage}% (Low Risk)")
                                elif percentage < 60:
                                    st.warning(f"📊 Plagiarism Risk: {percentage}% (Medium Risk)")
                                else:
                                    st.error(f"📊 Plagiarism Risk: {percentage}% (High Risk)")
                    except Exception as e:
                        st.error(f"Error parsing plagiarism percentage: {str(e)}")
                    
                    st.subheader("📝 Plagiarism Analysis:")
                    st.write(plagiarism_result)
                    
                    # Optionally save to DB if desired...
                    # if "current_student_id" in st.session_state and uploaded_file:
                    #     try:
                    #         conn = init_db()
                    #         c = conn.cursor()
                    #         c.execute('''
                    #         CREATE TABLE IF NOT EXISTS document_plagiarism_checks (
                    #             id INTEGER PRIMARY KEY AUTOINCREMENT,
                    #             student_id TEXT,
                    #             filename TEXT,
                    #             plagiarism_score REAL,
                    #             timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                    #         )
                    #         ''')
                    #         
                    #         # Attempt to store the extracted percentage
                    #         extracted_percentage = 0
                    #         try:
                    #             extracted_percentage = float(percentage_text.replace('%', '').strip())
                    #         except:
                    #             pass
                    # 
                    #         c.execute('''
                    #         INSERT INTO document_plagiarism_checks (student_id, filename, plagiarism_score, timestamp)
                    #         VALUES (?, ?, ?, ?)
                    #         ''', (
                    #             st.session_state.current_student_id,
                    #             uploaded_file.name,
                    #             extracted_percentage,
                    #             datetime.now()
                    #         ))
                    #         
                    #         conn.commit()
                    #         conn.close()
                    #     except Exception as e:
                    #         st.error(f"Error saving plagiarism check to database: {str(e)}")

            else:
                st.warning("⚠️ Please upload a valid document first.")

    # -- Local Reports Comparison --
    with col3:
        if st.button("🔎 Compare with Local Reports"):
            if student_submission.strip():
                with st.spinner("Comparing text with local PDF/DOCX reports..."):
                    results = compute_local_plagiarism_scores(student_submission, "local_reports")

                st.subheader("Local Similarity Scores (TF-IDF)")
                if "Error" in results and isinstance(results["Error"], str):
                    st.error(results["Error"])
                else:
                    # Sort by highest similarity first
                    items = sorted(
                        results.items(),
                        key=lambda x: x[1] if isinstance(x[1], float) else 0,
                        reverse=True
                    )
                    for fname, score in items:
                        if isinstance(score, float):
                            st.write(f"• **{fname}** → Similarity: {score:.2f}")
                            if score >= 0.8:
                                st.warning("High similarity! Potential plagiarism.")
                        else:
                            st.error(f"{fname}: {score}")
            else:
                st.warning("⚠️ Please upload a valid document first.")


