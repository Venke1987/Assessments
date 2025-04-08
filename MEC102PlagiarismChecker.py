import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import random
import openai
import json
from datetime import datetime
import fitz  # PyMuPDF for PDFs
import docx  # For Word document handling
import ast  # For checking Python syntax
import io  # For handling uploaded files
import sqlite3
import os
from io import BytesIO
import xgboost as xgb
import numpy as np
import cv2
import time
from PIL import Image, ImageDraw
from PIL import Image, ImageDraw
from io import BytesIO
import unicodedata
import re
# Add this near your other imports
from ai_tutor_helper import setup_ai_tutor, display_ai_tutor_page
# For Live OMR Test webcam functionality
try:
    from streamlit_webrtc import webrtc_streamer, VideoTransformerBase, RTCConfiguration
    import av
    import threading
    import queue
except ImportError:
    pass

# Uncomment and install these if needed
# !pip install opencv-python face_recognition xgboost scikit-learn streamlit_webrtc pillow av streamlit_webrtc

try:
    import cv2  # for Computer Vision (OMR, Face/Eye detection, etc.)
except ImportError:
    pass

try:
    import face_recognition  # for facial recognition
except ImportError:
    pass

try:
    import xgboost as xgb  # for XGBoost-based predictions
except ImportError:
    pass

try:
    from sklearn.ensemble import RandomForestClassifier
except ImportError:
    pass
from dotenv import load_dotenv
load_dotenv()

# ---------------- SET YOUR OPENAI API KEY HERE ----------------
openai.api_key = st.secrets["OPENAI_API_KEY"]
client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])  # Use Streamlit Secrets API Key

# --- Authentication ---
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

def authenticate_user():
    st.sidebar.markdown("🔐 **Faculty Login**")
    password = st.sidebar.text_input("Enter Password", type="password")
    if password == "admin@123":
        st.session_state.authenticated = True
        st.success("🔓 Access granted.")
    else:
        st.sidebar.warning("Invalid password.")

if not st.session_state.authenticated:
    authenticate_user()
    if not st.session_state.authenticated:
        st.stop()

# ---------------- DATABASE INIT ----------------
def init_db():
    conn = sqlite3.connect('student_quiz_history.db')
    c = conn.cursor()
    c.execute('''
    CREATE TABLE IF NOT EXISTS quiz_results (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        student_id TEXT NOT NULL,
        student_name TEXT NOT NULL,
        topic TEXT NOT NULL,
        score INTEGER NOT NULL,
        total_questions INTEGER NOT NULL,
        timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    conn.commit()
    return conn

def save_quiz_result(student_id, student_name, topic, score, total_questions):
    conn = init_db()
    c = conn.cursor()
    c.execute('''
    INSERT INTO quiz_results (student_id, student_name, topic, score, total_questions, timestamp)
    VALUES (?, ?, ?, ?, ?, ?)
    ''', (student_id, student_name, topic, score, total_questions, datetime.now()))
    conn.commit()
    conn.close()

def get_student_quiz_history(student_id=None):
    conn = init_db()
    c = conn.cursor()
    if student_id:
        c.execute('''
        SELECT student_id, student_name, topic, score, total_questions, timestamp
        FROM quiz_results
        WHERE student_id = ?
        ORDER BY timestamp DESC
        ''', (student_id,))
    else:
        c.execute('''
        SELECT student_id, student_name, topic, score, total_questions, timestamp
        FROM quiz_results
        ORDER BY timestamp DESC
        ''')
    results = c.fetchall()
    conn.close()
    if results:
        return pd.DataFrame(results, columns=['student_id', 'student_name', 'topic', 'score', 'total_questions', 'timestamp'])
    return pd.DataFrame()

def extract_text_from_docx(uploaded_file):
    """Extract text from a .docx file"""
    try:
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def extract_text_from_pdf(uploaded_file):
    """Extract text from a .pdf file."""
    try:
        text = ""
        with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text() + "\n"
        return text
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def check_python_syntax(code):
    """Check for syntax errors in uploaded Python code."""
    try:
        ast.parse(code)
        return None  # No syntax errors
    except SyntaxError as e:
        return f"Syntax Error: {e}"

# ---------------- LIVE OMR TEST FUNCTIONS ----------------
# Video processor for proctoring during the test
# Video processor for proctoring during the test
# Enhanced Video Processor for AI Proctoring
from streamlit_webrtc import VideoTransformerBase
class EnhancedVideoProcessor(VideoTransformerBase):
    def __init__(self):
        # Face and eye detection using Haar cascades
        self.face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        self.eye_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_eye.xml')
        
        # State tracking
        self.frame_count = 0
        self.last_warning_time = time.time() - 10  # Initialize to allow immediate warnings
        self.face_detected = False
        self.eyes_detected = False
        self.multiple_faces = False
        self.looking_away_duration = 0
        self.suspicious_movements = 0
        
        # Audio monitoring
        self.audio_threshold = 0.1
        self.last_audio_warning_time = time.time() - 10
        self.noise_detection_window = []
        
        # Warning management
        self.should_warn = False
        self.warning_message = ""
        self.warning_type = None
        
        # Debug and analytics
        self.debug_info = {
            "frames_processed": 0,
            "face_detections": 0,
            "eye_detections": 0,
            "warnings_triggered": 0,
            "audio_warnings": 0,
            "multi_face_warnings": 0
        }
        
        # Movement detection
        self.prev_frame = None
        self.movement_threshold = 0.05
        self.movement_detection_counter = 0
        
    def detect_suspicious_movement(self, current_frame):
        """Detect rapid or unusual movements in the frame"""
        if self.prev_frame is None:
            self.prev_frame = current_frame
            return False
        
        # Calculate frame difference
        frame_diff = cv2.absdiff(current_frame, self.prev_frame)
        gray_diff = cv2.cvtColor(frame_diff, cv2.COLOR_BGR2GRAY)
        _, thresh_diff = cv2.threshold(gray_diff, 30, 255, cv2.THRESH_BINARY)
        
        # Calculate movement percentage
        movement_percentage = np.sum(thresh_diff) / (thresh_diff.shape[0] * thresh_diff.shape[1] * 255)
        
        # Update previous frame
        self.prev_frame = current_frame
        
        # Check if movement exceeds threshold
        if movement_percentage > self.movement_threshold:
            self.movement_detection_counter += 1
            # Only trigger warning if we have detected excessive movement for several frames
            if self.movement_detection_counter > 5:
                return True
        else:
            self.movement_detection_counter = max(0, self.movement_detection_counter - 1)
            
        return False

    def transform(self, frame):
        img = frame.to_ndarray(format="bgr24")
        self.frame_count += 1
        self.debug_info["frames_processed"] += 1
        
        # Process every frame for better responsiveness
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # More sensitive face detection
        faces = self.face_cascade.detectMultiScale(
            gray, 
            scaleFactor=1.05,
            minNeighbors=3,
            minSize=(20, 20),
            flags=cv2.CASCADE_SCALE_IMAGE
        )

        # Update face detection state
        prev_face_detected = self.face_detected
        self.face_detected = len(faces) > 0
        
        # Check for multiple faces - potential cheating
        self.multiple_faces = len(faces) > 1
        
        if self.face_detected:
            self.debug_info["face_detections"] += 1
        
        current_time = time.time()

        # Check for suspicious movements
        suspicious_movement = self.detect_suspicious_movement(img)
        if suspicious_movement and (current_time - self.last_warning_time) > 3:
            self.last_warning_time = current_time
            self.should_warn = True
            self.warning_type = "movement"
            self.warning_message = "⚠️ Warning: Excessive movement detected. Please remain still during the test."
            if "warnings_count" in st.session_state:
                st.session_state.warnings_count += 1

        # If multiple faces detected - possible cheating
        if self.multiple_faces and (current_time - self.last_warning_time) > 3:
            self.last_warning_time = current_time
            self.debug_info["multi_face_warnings"] += 1
            self.should_warn = True
            self.warning_type = "multiple_faces"
            self.warning_message = "⚠️ Warning: Multiple faces detected. Only you should be visible during the test."
            if "warnings_count" in st.session_state:
                st.session_state.warnings_count += 1

        # If face is detected
        if self.face_detected:
            # Update Streamlit session state if it exists
            if "last_detection_time" in st.session_state:
                st.session_state.last_detection_time = current_time
            if "looking_away_start" in st.session_state:
                st.session_state.looking_away_start = None

            # Draw rectangle around faces and check for eyes
            for (x, y, w, h) in faces:
                cv2.rectangle(img, (x, y), (x+w, y+h), (255, 0, 0), 2)
                roi_gray = gray[y:y+h, x:x+w]
                roi_color = img[y:y+h, x:x+w]
                
                # More sensitive eye detection
                eyes = self.eye_cascade.detectMultiScale(
                    roi_gray,
                    scaleFactor=1.05,
                    minNeighbors=2,
                    minSize=(5, 5)
                )
                
                prev_eyes_detected = self.eyes_detected
                self.eyes_detected = len(eyes) > 0
                
                if self.eyes_detected:
                    self.debug_info["eye_detections"] += 1
                
                for (ex, ey, ew, eh) in eyes:
                    cv2.rectangle(roi_color, (ex, ey), (ex+ew, ey+eh), (0, 255, 0), 2)
                
                # If eyes not detected, trigger warning
                if not self.eyes_detected and (current_time - self.last_warning_time) > 2:
                    self.last_warning_time = current_time
                    self.debug_info["warnings_triggered"] += 1
                    
                    # Set direct warning flag
                    self.should_warn = True
                    self.warning_type = "eyes"
                    self.warning_message = "⚠️ Warning: Eyes not detected. Please look at the screen."
                    
                    # Update session state
                    if "warnings_count" in st.session_state:
                        st.session_state.warnings_count += 1
        else:
            # Face not detected - update looking away time
            if "looking_away_start" in st.session_state:
                if st.session_state.looking_away_start is None:
                    st.session_state.looking_away_start = current_time
                # Reduced time threshold for warning (1.5 seconds)
                elif (current_time - st.session_state.looking_away_start) > 1.5 and (current_time - self.last_warning_time) > 2:
                    self.last_warning_time = current_time
                    self.debug_info["warnings_triggered"] += 1
                    
                    # Set direct warning flag
                    self.should_warn = True
                    self.warning_type = "face"
                    self.warning_message = "⚠️ Warning: Face not detected. Please stay in front of the camera."
                    
                    # Update session state
                    if "warnings_count" in st.session_state:
                        st.session_state.warnings_count += 1

        # Check if warnings threshold reached - terminate test
        if "warnings_count" in st.session_state and "test_in_progress" in st.session_state:
            if st.session_state.warnings_count >= 3 and st.session_state.test_in_progress:
                st.session_state.test_in_progress = False
                st.session_state.test_terminated = True
                st.session_state.termination_reason = "Multiple violations of test integrity rules"

        # Add status information to the frame
        status_text = "Status: Face Detected ✓" if self.face_detected else "Status: Face Not Detected ✗"
        cv2.putText(img, status_text, (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 0.7, 
                   (0, 255, 0) if self.face_detected else (0, 0, 255), 2)
        
        eye_status = "Eyes Detected ✓" if self.eyes_detected else "Eyes Not Detected ✗"
        cv2.putText(img, eye_status, (10, 60), cv2.FONT_HERSHEY_SIMPLEX, 0.7, 
                   (0, 255, 0) if self.eyes_detected else (0, 0, 255), 2)
        
        # Add warning count
        if "warnings_count" in st.session_state:
            warning_color = (0, 255, 0)  # Green for 0 warnings
            if st.session_state.warnings_count == 1:
                warning_color = (0, 255, 255)  # Yellow for 1 warning
            elif st.session_state.warnings_count == 2:
                warning_color = (0, 165, 255)  # Orange for 2 warnings
            elif st.session_state.warnings_count >= 3:
                warning_color = (0, 0, 255)    # Red for 3+ warnings
                
            warnings_text = f"Warnings: {st.session_state.warnings_count}/3"
            cv2.putText(img, warnings_text, (10, 90), cv2.FONT_HERSHEY_SIMPLEX, 0.7, warning_color, 2)
        
        # Add debug info
        debug_text = f"Processed: {self.debug_info['frames_processed']} | Warnings: {self.debug_info['warnings_triggered']}"
        cv2.putText(img, debug_text, (10, img.shape[0] - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (255, 255, 255), 1)

        return img

    def recv(self, frame):
        # Process video
        if frame.format.name == "video":
            img = self.transform(frame)
            
            # Direct warning mechanism
            if self.should_warn:
                # Apply warning to session state
                if "latest_warning" in st.session_state:
                    st.session_state.latest_warning = self.warning_message
                if "latest_warning_time" in st.session_state:
                    st.session_state.latest_warning_time = time.time()
                if "latest_warning_type" in st.session_state:
                    st.session_state.latest_warning_type = self.warning_type
                    
                self.should_warn = False  # Reset flag
            
            return av.VideoFrame.from_ndarray(img, format="bgr24")
        
        # Process audio
        elif frame.format.name == "audio":
            sound = frame.to_ndarray()
            sound_level = np.abs(sound).max()
            self.noise_detection_window.append(sound_level)
            
            # Keep a rolling window of sound levels
            if len(self.noise_detection_window) > 20:
                self.noise_detection_window.pop(0)
            
            # Check for sustained high audio levels - possible talking
            if len(self.noise_detection_window) > 5:
                current_time = time.time()
                average_level = np.mean(self.noise_detection_window)
                
                if average_level > self.audio_threshold and current_time - self.last_audio_warning_time > 5:
                    self.last_audio_warning_time = current_time
                    self.debug_info["audio_warnings"] += 1
                    
                    if "warnings_count" in st.session_state:
                        st.session_state.warnings_count += 1
                        
                    # Update session state with audio warning
                    if "latest_warning" in st.session_state:
                        st.session_state.latest_warning = "⚠️ Warning: High audio levels detected. Please remain quiet during the test."
                    if "latest_warning_time" in st.session_state:
                        st.session_state.latest_warning_time = time.time()
                    if "latest_warning_type" in st.session_state:
                        st.session_state.latest_warning_type = "audio"
            
            return frame
        
        return frame
# Function to initialize proctoring session state
def init_proctoring_state():
    """Initialize all session state variables needed for proctoring"""
    if "warnings_count" not in st.session_state:
        st.session_state.warnings_count = 0
    if "looking_away_start" not in st.session_state:
        st.session_state.looking_away_start = None
    if "last_detection_time" not in st.session_state:
        st.session_state.last_detection_time = time.time()
    if "latest_warning" not in st.session_state:
        st.session_state.latest_warning = None
    if "latest_warning_time" not in st.session_state:
        st.session_state.latest_warning_time = None
    if "latest_warning_type" not in st.session_state:
        st.session_state.latest_warning_type = None
    if "test_terminated" not in st.session_state:
        st.session_state.test_terminated = False
    if "termination_reason" not in st.session_state:
        st.session_state.termination_reason = None

# Function to display and manage proctoring warnings
def display_proctoring_warnings():
    """Display proctoring warnings and handle test termination"""
    # Create container for warnings
    warning_container = st.empty()
    
    # Display current warning if it exists and is recent (within last 10 seconds)
    current_time = time.time()
    if (st.session_state.latest_warning_time is not None and 
            current_time - st.session_state.latest_warning_time < 10):
        
        warning_type = st.session_state.latest_warning_type
        warning_message = st.session_state.latest_warning
        
        if warning_type == "multiple_faces":
            warning_container.error(warning_message)
        else:
            warning_container.warning(warning_message)
            
        # Also display warning in sidebar for higher visibility
        with st.sidebar:
            if warning_type == "multiple_faces":
                st.error(warning_message)
            else:
                st.warning(warning_message)
    
    # Display warning count prominently
    current_warnings = st.session_state.warnings_count
    col1, col2 = st.columns([3, 1])
    
    with col1:
        if current_warnings > 0:
            st.markdown(f"### ⚠️ Current warnings: {current_warnings}/3")
        else:
            st.markdown(f"### Current warnings: {current_warnings}/3")
    
    with col2:
        if st.button("Reset Warnings"):
            st.session_state.warnings_count = 0
            st.session_state.latest_warning = None
            st.session_state.latest_warning_time = None
            st.session_state.latest_warning_type = None
            st.rerun()
    
    # Use a progress bar to visualize warnings with color coding
    if current_warnings == 0:
        st.progress(current_warnings/3, "Good standing")
    elif current_warnings == 1:
        st.warning("First warning issued")
        st.progress(current_warnings/3)
    elif current_warnings == 2:
        st.warning("⚠️⚠️ Second warning issued! One more warning will terminate the test.")
        st.progress(current_warnings/3)
    elif current_warnings >= 3:
        st.error("❌ Maximum warnings reached")
        st.progress(1.0)
    
    # Handle test termination if too many warnings
    if st.session_state.test_terminated:
        st.error(f"❌ Test terminated: {st.session_state.termination_reason}")
        if st.button("Return to Test Setup"):
            st.session_state.test_in_progress = False
            st.session_state.test_completed = False
            st.session_state.test_terminated = False
            st.session_state.warnings_count = 0
            st.rerun()
# Function to generate questions using OpenAI's API
def generate_questions(topic, num_questions):
    """Generate multiple choice questions on the given topic using ChatGPT API"""
    prompt = f"""Generate a multiple-choice quiz with {num_questions} questions on the topic '{topic}'. 
    Each question should have 4 options labeled A, B, C, D and indicate the correct answer. 
    Format your response as a JSON array of objects with these exact keys:
    - "question": the question text
    - "options": an array of 4 strings representing the options
    - "correct": the correct answer (A, B, C, or D)
    - "explanation": brief explanation of the correct answer
    
    The JSON must be valid and properly formatted."""
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an AI that generates quiz questions. Return valid JSON only with no other text."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        response_text = response.choices[0].message.content.strip()
        
        # Extract JSON if it's wrapped in code blocks
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0].strip()
            
        questions_data = json.loads(response_text)
        return questions_data
    
    except Exception as e:
        st.error(f"Error generating questions: {str(e)}")
        st.write("Response received:", response_text if 'response_text' in locals() else "No response")
        return None
# Function to create an interactive OMR sheet
def create_omr_sheet(num_questions, current_answers=None):
    """Create an interactive OMR sheet with the given number of questions"""
    if current_answers is None:
        current_answers = [None] * num_questions
        
    # Create a new image
    width, height = 600, 100 + (num_questions * 50)
    img = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(img)
    
    # Draw title
    draw.text((width//2 - 80, 20), "OMR ANSWER SHEET", fill='black')
    draw.line([(50, 50), (width-50, 50)], fill='black', width=2)
    
    # Draw question numbers and option bubbles
    for i in range(num_questions):
        y_pos = 80 + (i * 50)
        # Question number
        draw.text((30, y_pos+5), f"{i+1}.", fill='black')
        
        # Option bubbles
        options = ['A', 'B', 'C', 'D']
        for j, option in enumerate(options):
            x_pos = 80 + (j * 100)
            # Draw option label
            draw.text((x_pos-5, y_pos-15), option, fill='black')
            
            # Draw circle
            radius = 15
            if current_answers[i] == option:
                # Filled circle if selected
                draw.ellipse((x_pos-radius, y_pos-radius, x_pos+radius, y_pos+radius), outline='black', fill='black')
            else:
                # Empty circle if not selected
                draw.ellipse((x_pos-radius, y_pos-radius, x_pos+radius, y_pos+radius), outline='black')
    
    # Convert to bytes for displaying in Streamlit
    buf = BytesIO()
    img.save(buf, format='PNG')
    byte_im = buf.getvalue()
    
    return byte_im
# Function to evaluate the test and provide feedback
def evaluate_test(questions, student_answers):
    """Evaluate the test and get feedback using ChatGPT"""
    correct_answers = [q["correct"] for q in questions]
    score = sum(1 for i in range(len(correct_answers)) if student_answers[i] == correct_answers[i])
    
    feedback_prompt = f"""Evaluate the following test results:
    
    Questions and correct answers:
    """
    
    for idx, q in enumerate(questions, start=1):
        feedback_prompt += f"Q{idx}: {q['question']}\n"
        feedback_prompt += f"Options: A: {q['options'][0]}, B: {q['options'][1]}, C: {q['options'][2]}, D: {q['options'][3]}\n"
        feedback_prompt += f"Correct Answer: {q['correct']}\n"
        feedback_prompt += f"Student's Answer: {student_answers[idx-1] if idx-1 < len(student_answers) else 'Not answered'}\n\n"
    
    feedback_prompt += f"""
    Student scored {score} out of {len(questions)}.
    
    Please provide:
    1. Overall assessment of the student's performance
    2. Specific feedback on incorrect answers
    3. Areas for improvement
    4. Study recommendations
    
    Use a supportive and constructive tone.
    """
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an AI that provides educational feedback and assessment."},
                {"role": "user", "content": feedback_prompt}
            ],
            temperature=0.4,
            max_tokens=800
        )
        
        feedback = response.choices[0].message.content.strip()
        return {
            "score": score,
            "total": len(questions),
            "percentage": round((score / len(questions)) * 100, 1),
            "feedback": feedback
        }
    except Exception as e:
        st.error(f"Error generating feedback: {str(e)}")
        return {
            "score": score,
            "total": len(questions),
            "percentage": round((score / len(questions)) * 100, 1),
            "feedback": "Error generating detailed feedback."
        }
# Add this function to generate questions from uploaded file content
def generate_questions_from_file(file_content, num_questions):
    """Generate multiple choice questions based on uploaded file content using ChatGPT API"""
    prompt = f"""Based on the following document content, generate a multiple-choice quiz with {num_questions} questions. 
    Each question should have 4 options labeled A, B, C, D and indicate the correct answer. 
    Format your response as a JSON array of objects with these exact keys:
    - "question": the question text
    - "options": an array of 4 strings representing the options
    - "correct": the correct answer (A, B, C, or D)
    - "explanation": brief explanation of the correct answer
    
    The JSON must be valid and properly formatted.
    
    Document Content:
    {file_content}
    """
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an AI that generates quiz questions from documents. Return valid JSON only with no other text."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        response_text = response.choices[0].message.content.strip()
        
        # Extract JSON if it's wrapped in code blocks
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0].strip()
            
        questions_data = json.loads(response_text)
        return questions_data
    
    except Exception as e:
        st.error(f"Error generating questions: {str(e)}")
        st.write("Response received:", response_text if 'response_text' in locals() else "No response")
        return None

# Modify the display_live_omr_test function to include file upload
def display_live_omr_test():
    # Add at the beginning of display_live_omr_test function
    st.title("📝 Interactive Live OMR Test with AI Proctoring")

    # Auto-refresh mechanism
    if "last_refresh_time" not in st.session_state:
        st.session_state.last_refresh_time = time.time()

    current_time = time.time()
    # Force a rerun every 3 seconds to check for new warnings
    if current_time - st.session_state.last_refresh_time > 3:
        st.session_state.last_refresh_time = current_time
        try:
            st.experimental_rerun()  # Use experimental_rerun for compatibility
        except:
            # Fallback to rerun if available
            if hasattr(st, 'rerun'):
                st.rerun()
   
    # Initialize proctoring state
    init_proctoring_state()
    
    # Create a placeholder for warnings
    warning_placeholder = st.empty()
    
    # Set up layout for test and proctoring
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📹 Proctoring Camera")
        
        # WebRTC configuration for proctoring with improved settings
        rtc_configuration = RTCConfiguration(
            {"iceServers": [{"urls": ["stun:stun.l.google.com:19302"]}]}
        )
        
        # Start webcam with enhanced processor and better configuration
        try:
            webrtc_ctx = webrtc_streamer(
                key="enhanced_proctoring",
                video_processor_factory=EnhancedVideoProcessor,
                rtc_configuration=rtc_configuration,
                media_stream_constraints={
                    "video": True,
                    "audio": True
                },
                async_processing=True,
                video_html_attrs={
                    "style": {"width": "100%", "height": "auto", "margin": "0 auto"},
                    "controls": False,
                    "autoPlay": True,
                },
            )
        except Exception as e:
            st.error(f"Error initializing webcam: {str(e)}")
            webrtc_ctx = None
        
        # Check if camera is connected before showing proctoring
        if webrtc_ctx and webrtc_ctx.state and webrtc_ctx.state.playing:
            # Display proctoring warnings and status
            display_proctoring_warnings()
            
            # Add a debug section only visible with a toggle
            with st.expander("Debug Information (Click to expand)"):
                st.write(f"Face detected: {getattr(webrtc_ctx.video_processor, 'face_detected', 'N/A') if webrtc_ctx and webrtc_ctx.video_processor else 'No webcam'}")
                st.write(f"Eyes detected: {getattr(webrtc_ctx.video_processor, 'eyes_detected', 'N/A') if webrtc_ctx and webrtc_ctx.video_processor else 'No webcam'}")
                st.write(f"Multiple faces: {getattr(webrtc_ctx.video_processor, 'multiple_faces', 'N/A') if webrtc_ctx and webrtc_ctx.video_processor else 'No webcam'}")
                st.write(f"Warning count: {st.session_state.warnings_count}")
                st.write(f"Latest warning: {st.session_state.latest_warning}")
                st.write(f"Warning time: {st.session_state.latest_warning_time}")
                st.write(f"Warning type: {st.session_state.latest_warning_type}")
            
                # Add a manual warning button for testing
                if st.button("Trigger Test Warning"):
                    st.session_state.warnings_count += 1
                    st.session_state.latest_warning = "This is a test warning"
                    st.session_state.latest_warning_time = time.time()
                    st.session_state.latest_warning_type = "test"
                    try:
                        st.experimental_rerun()
                    except:
                        if hasattr(st, 'rerun'):
                            st.rerun()
        else:
            st.warning("⚠️ Camera not connected. Please ensure your camera is working and you've granted browser permissions.")
            st.info("If you're having trouble connecting, try refreshing the page or using a different browser.")
            
            # Add alternative capture method
            if st.button("Use Alternative Camera Method"):
                st.session_state.use_alternative_camera = True
                try:
                    st.experimental_rerun()
                except:
                    if hasattr(st, 'rerun'):
                        st.rerun()
                
            if st.session_state.get("use_alternative_camera", False):
                st.success("Using alternative camera method")
                # Create a placeholder with simulated camera feed
                img = Image.new('RGB', (640, 480), color='white')
                draw = ImageDraw.Draw(img)
                draw.text((320, 240), "Camera Simulation Active", fill='black')
                st.image(img, channels="RGB", use_container_width=True)
                
                # Add test warning controls for the alternative mode
                if st.button("Simulate Warning"):
                    st.session_state.warnings_count = min(st.session_state.warnings_count + 1, 3)
                    st.session_state.latest_warning = "Simulated Warning"
                    st.session_state.latest_warning_time = time.time()
                    st.session_state.latest_warning_type = "test" 
                    try:
                        st.experimental_rerun()
                    except:
                        if hasattr(st, 'rerun'):
                            st.rerun()
        
        # Proctoring rules
        st.caption("📝 Proctoring Rules:")
        st.caption("1. Keep your face visible at all times")
        st.caption("2. Look at the screen (eyes must be visible)")
        st.caption("3. No other people in frame")
        st.caption("4. Remain quiet during the test")
        st.caption("5. Avoid excessive movement")
        st.caption("6. Three warnings will terminate the test")
    
    # If test is terminated, don't show test content
    if st.session_state.test_terminated:
        return
        
    # Step 1: Configure the test if not already in progress
    if not st.session_state.test_in_progress and not st.session_state.test_completed:
        with col2:
            st.header("Configure Your Test")
            
            # Student selection
            student_id = st.selectbox("Select your Student ID", list(students_data.keys()))
            st.write(f"Student Name: {students_data[student_id]['name']}")
            st.session_state.current_student_id = student_id
            st.session_state.current_student_name = students_data[student_id]["name"]
            
            # NEW: Add file upload for generating questions
            st.subheader("Upload Content for Questions")
            uploaded_file = st.file_uploader("Upload a file (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"])
            
            file_content = None
            if uploaded_file is not None:
                st.success(f"File uploaded: {uploaded_file.name}")
                
                # Extract text from different file types
                try:
                    if uploaded_file.name.endswith('.txt'):
                        file_content = uploaded_file.read().decode('utf-8')
                    elif uploaded_file.name.endswith('.pdf'):
                        file_content = extract_text_from_pdf(uploaded_file)
                    elif uploaded_file.name.endswith('.docx'):
                        file_content = extract_text_from_docx(uploaded_file)
                    
                    # Show preview of extracted content
                    with st.expander("Content Preview"):
                        st.write(file_content[:500] + "..." if len(file_content) > 500 else file_content)
                except Exception as e:
                    st.error(f"Error extracting content from file: {str(e)}")
            
            # Allow backup topic entry if no file is uploaded
            if not file_content:
                test_topic = st.text_input("Enter test topic (if no file uploaded):", "Computer Science Fundamentals")
                st.session_state.current_test_topic = test_topic
            
            num_questions = st.number_input("Number of questions:", min_value=1, max_value=20, value=5)
            
            # Global secure OpenAI key setting
                        
            # Add a camera check before starting the test
            st.subheader("Camera & Microphone Check")
            st.write("Please ensure your camera and microphone are working correctly.")
            
            # Check if webcam is active and running or we're using alternative mode
            camera_ready = (webrtc_ctx and webrtc_ctx.state and webrtc_ctx.state.playing) or st.session_state.get("use_alternative_camera", False)
            
            if camera_ready:
                st.success("✅ Camera and microphone are working!")
                
                # Start test button
                if st.button("Start Test"):
                    with st.spinner("Generating test questions..."):
                        if file_content:
                            # Generate questions from file content
                            questions = generate_questions_from_file(file_content, num_questions)
                        else:
                            # Fallback to topic-based question generation
                            questions = generate_questions(test_topic, num_questions)
                    
                    if questions:
                        st.session_state.live_omr_questions = questions
                        st.session_state.student_answers = [None] * len(questions)
                        st.session_state.current_question_index = 0
                        st.session_state.warnings_count = 0  # Reset warnings
                        st.session_state.test_in_progress = True
                        st.session_state.test_completed = False
                        st.session_state.test_terminated = False
                        try:
                            st.experimental_rerun()
                        except:
                            if hasattr(st, 'rerun'):
                                st.rerun()
                    else:
                        st.error("Failed to generate questions. Please try again.")
            else:
                st.error("⚠️ Camera access is required. Please allow camera access and try again.")
                
    # Step 2: Display the proctored test
    elif st.session_state.test_in_progress:
        with col2:
            st.header("🔍 OMR Test in Progress")
            
            if st.session_state.live_omr_questions:
                current_q = st.session_state.live_omr_questions[st.session_state.current_question_index]
                st.subheader(f"Question {st.session_state.current_question_index + 1} of {len(st.session_state.live_omr_questions)}")
                st.write(current_q["question"])
                
                # Display options
                for i, option in enumerate(['A', 'B', 'C', 'D']):
                    st.write(f"{option}: {current_q['options'][i]}")
            
            # Display interactive OMR sheet
            st.subheader("📑 Answer Sheet")
            st.caption("Select your answer below")
            
            # Create OMR sheet image
            omr_image = create_omr_sheet(len(st.session_state.live_omr_questions), st.session_state.student_answers)
            
            # Display the image
            omr_placeholder = st.empty()
            omr_placeholder.image(omr_image, use_container_width=True)
            
            # Handle test navigation and answers
            nav_col1, nav_col2, nav_col3 = st.columns([1, 2, 1])
            
            with nav_col1:
                if st.button("⬅️ Previous") and st.session_state.current_question_index > 0:
                    st.session_state.current_question_index -= 1
                    try:
                        st.experimental_rerun()
                    except:
                        if hasattr(st, 'rerun'):
                            st.rerun()
            
            with nav_col2:
                # Allow selecting answer via buttons
                selected_option = st.selectbox(
                    "Select answer:", 
                    [None, 'A', 'B', 'C', 'D'], 
                    index=0 if st.session_state.student_answers[st.session_state.current_question_index] is None else 
                         ['A', 'B', 'C', 'D'].index(st.session_state.student_answers[st.session_state.current_question_index]) + 1
                )
                
                if selected_option:
                    st.session_state.student_answers[st.session_state.current_question_index] = selected_option
                    # Update the OMR sheet
                    omr_image = create_omr_sheet(len(st.session_state.live_omr_questions), st.session_state.student_answers)
                    omr_placeholder.image(omr_image, use_container_width=True)
            
            with nav_col3:
                if st.button("Next ➡️") and st.session_state.current_question_index < len(st.session_state.live_omr_questions) - 1:
                    st.session_state.current_question_index += 1
                    try:
                        st.experimental_rerun()
                    except:
                        if hasattr(st, 'rerun'):
                            st.rerun()
            
            # Submit button
            st.markdown("---")
            if st.button("📝 Submit Test", use_container_width=True):
                # Check if all questions have been answered
                if None in st.session_state.student_answers:
                    st.warning("Please answer all questions before submitting.")
                else:
                    st.session_state.test_in_progress = False
                    st.session_state.test_completed = True
                    try:
                        st.experimental_rerun()
                    except:
                        if hasattr(st, 'rerun'):
                            st.rerun()
    
    # Step 3: Show test results and feedback
    elif st.session_state.test_completed:
        with col2:
            st.header("📊 Test Results")
            
            # Display final OMR sheet
            st.subheader("Your Completed OMR Sheet")
            final_omr = create_omr_sheet(len(st.session_state.live_omr_questions), st.session_state.student_answers)
            st.image(final_omr, use_container_width=True)
            
            # Evaluate the test
            with st.spinner("Evaluating your test..."):
                evaluation = evaluate_test(st.session_state.live_omr_questions, st.session_state.student_answers)
            
            # Display score
            st.subheader(f"Your Score: {evaluation['score']}/{evaluation['total']} ({evaluation['percentage']}%)")
            
            # Display progress bar
            score_percent = evaluation['percentage'] / 100
            st.progress(score_percent)
            
            # Add celebrations for good scores
            if score_percent >= 0.8:
                st.balloons()
                st.success("🎉 Excellent job!")
            elif score_percent >= 0.6:
                st.success("👍 Good work!")
            else:
                st.info("📚 Keep practicing!")
            
            # Display detailed feedback
            with st.expander("Detailed Feedback", expanded=True):
                st.write(evaluation['feedback'])
            
            # Question-by-question review
            st.subheader("Question-by-Question Review")
            for i, question in enumerate(st.session_state.live_omr_questions):
                correct = st.session_state.student_answers[i] == question['correct']
                expander_title = f"Q{i+1}: {'✅ Correct' if correct else '❌ Incorrect'}"
                with st.expander(expander_title):
                    st.write(f"**Question:** {question['question']}")
                    st.write(f"**Your answer:** {st.session_state.student_answers[i]} - {question['options'][ord(st.session_state.student_answers[i]) - ord('A')]}")
                    st.write(f"**Correct answer:** {question['correct']} - {question['options'][ord(question['correct']) - ord('A')]}")
                    st.write(f"**Explanation:** {question['explanation']}")
            
            # Start a new test button
            if st.button("Start a New Test", use_container_width=True):
                st.session_state.test_in_progress = False
                st.session_state.test_completed = False
                st.session_state.test_terminated = False
                st.session_state.live_omr_questions = None
                st.session_state.student_answers = []
                st.session_state.current_question_index = 0
                st.session_state.warnings_count = 0
                try:
                    st.experimental_rerun()
                except:
                    if hasattr(st, 'rerun'):
                        st.rerun()

# ---------------- SAMPLE STUDENT DATA ----------------
students_data = {
    "SEEE001": {"name": "Adhithya V", "proficiency": {"Fuzzy Logic": {1: 0.59, 2: 0.64, 3: 0.69, 4: 0.74, 5: 0.79, 6: 0.84}}},
    "SEEE002": {"name": "Akety Manjunath", "proficiency": {"Data Science": {1: 0.6, 2: 0.65, 3: 0.7, 4: 0.75, 5: 0.8, 6: 0.85}}},
    "SEEE003": {"name": "Aravind S", "proficiency": {"Bayesian Learning": {1: 0.37, 2: 0.42, 3: 0.47, 4: 0.52, 5: 0.57, 6: 0.62}}},
    "SEEE004": {"name": "Aswin S", "proficiency": {"NLP": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}},
    "SEEE005": {"name": "Avinash Kumar R", "proficiency": {"Data Science": {1: 0.69, 2: 0.74, 3: 0.79, 4: 0.84, 5: 0.89, 6: 0.94}}},
    "SEEE007": {"name": "Bhavya Sri B", "proficiency": {"IoT": {1: 0.59, 2: 0.64, 3: 0.69, 4: 0.74, 5: 0.79, 6: 0.84}}},
    "SEEE008": {"name": "Challagandla Anantha Pavan", "proficiency": {"Cybersecurity": {1: 0.7, 2: 0.75, 3: 0.8, 4: 0.85, 5: 0.9, 6: 0.95}}},
    "SEEE009": {"name": "Nagadarahas Kumar C S", "proficiency": {"Cybersecurity": {1: 0.6, 2: 0.65, 3: 0.7, 4: 0.75, 5: 0.8, 6: 0.85}}},
    "SEEE010": {"name": "Dodda Sri Pujitha", "proficiency": {"Edge AI": {1: 0.73, 2: 0.78, 3: 0.83, 4: 0.88, 5: 0.93, 6: 0.98}}},
    "SEEE011": {"name": "Dondapati Vallapala Yami", "proficiency": {"Deep Learning": {1: 0.58, 2: 0.63, 3: 0.68, 4: 0.73, 5: 0.78, 6: 0.83}}},
    "SEEE012": {"name": "Guduru Venkata Sai Karthikeya", "proficiency": {"Generative AI": {1: 0.42, 2: 0.47, 3: 0.52, 4: 0.57, 5: 0.62, 6: 0.67}}},
    "SEEE013": {"name": "Hamsitha P", "proficiency": {"AI Ethics": {1: 0.39, 2: 0.44, 3: 0.49, 4: 0.54, 5: 0.59, 6: 0.64}}},
    "SEEE014": {"name": "Harini M", "proficiency": {"Generative AI": {1: 0.5, 2: 0.55, 3: 0.6, 4: 0.65, 5: 0.7, 6: 0.75}}},
    "SEEE015": {"name": "Harrish B", "proficiency": {"Robotics": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE016": {"name": "Jivan Prasath S", "proficiency": {"Swarm Intelligence": {1: 0.51, 2: 0.56, 3: 0.61, 4: 0.66, 5: 0.71, 6: 0.76}}},
    "SEEE017": {"name": "Jonnalagadda Susrith", "proficiency": {"Quantum Computing": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE018": {"name": "Kamaleshwar M", "proficiency": {"Data Science": {1: 0.72, 2: 0.77, 3: 0.82, 4: 0.87, 5: 0.92, 6: 0.97}}},
    "SEEE019": {"name": "Karthik Periyakarupphan M", "proficiency": {"NLP": {1: 0.36, 2: 0.41, 3: 0.46, 4: 0.51, 5: 0.56, 6: 0.61}}},
    "SEEE020": {"name": "Kathirazagan V", "proficiency": {"Bayesian Learning": {1: 0.44, 2: 0.49, 3: 0.54, 4: 0.59, 5: 0.64, 6: 0.69}}},
    "SEEE021": {"name": "Kishore R", "proficiency": {"Robotics": {1: 0.72, 2: 0.77, 3: 0.82, 4: 0.87, 5: 0.92, 6: 0.97}}},
    "SEEE022": {"name": "Krishnakumar Aditi J", "proficiency": {"Computer Vision": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE023": {"name": "Logavarshini K", "proficiency": {"Optimization": {1: 0.7, 2: 0.75, 3: 0.8, 4: 0.85, 5: 0.9, 6: 0.95}}},
    "SEEE024": {"name": "Malapalli Charitha Reddy", "proficiency": {"Deep Learning": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}},
    "SEEE025": {"name": "Manas M Nair", "proficiency": {"Bayesian Learning": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE026": {"name": "Mervath M", "proficiency": {"Robotics": {1: 0.55, 2: 0.6, 3: 0.65, 4: 0.7, 5: 0.75, 6: 0.8}}},
    "SEEE027": {"name": "Nidhish Kumar K", "proficiency": {"Bayesian Learning": {1: 0.37, 2: 0.42, 3: 0.47, 4: 0.52, 5: 0.57, 6: 0.62}}},
    "SEEE028": {"name": "Palapati Jahnavi", "proficiency": {"Reinforcement Learning": {1: 0.65, 2: 0.7, 3: 0.75, 4: 0.8, 5: 0.85, 6: 0.9}}},
    "SEEE029": {"name": "Pola Kaarthi", "proficiency": {"Fuzzy Logic": {1: 0.74, 2: 0.79, 3: 0.84, 4: 0.89, 5: 0.94, 6: 0.99}}},
    "SEEE030": {"name": "Pullela Vaishnavi", "proficiency": {"Data Science": {1: 0.39, 2: 0.44, 3: 0.49, 4: 0.54, 5: 0.59, 6: 0.64}}},
    "SEEE031": {"name": "Raghul T", "proficiency": {"Generative AI": {1: 0.43, 2: 0.48, 3: 0.53, 4: 0.58, 5: 0.63, 6: 0.68}}},
    "SEEE032": {"name": "Rajaprabu K", "proficiency": {"Big Data": {1: 0.41, 2: 0.46, 3: 0.51, 4: 0.56, 5: 0.61, 6: 0.66}}},
    "SEEE033": {"name": "Rishi A", "proficiency": {"NLP": {1: 0.41, 2: 0.46, 3: 0.51, 4: 0.56, 5: 0.61, 6: 0.66}}},
    "SEEE034": {"name": "Rithick Reddy S", "proficiency": {"Swarm Intelligence": {1: 0.48, 2: 0.53, 3: 0.58, 4: 0.63, 5: 0.68, 6: 0.73}}},
    "SEEE035": {"name": "Tera Sai Tejeshwar Reddy", "proficiency": {"Swarm Intelligence": {1: 0.59, 2: 0.64, 3: 0.69, 4: 0.74, 5: 0.79, 6: 0.84}}},
    "SEEE036": {"name": "Saikirthiga R", "proficiency": {"Data Science": {1: 0.55, 2: 0.6, 3: 0.65, 4: 0.7, 5: 0.75, 6: 0.8}}},
    "SEEE037": {"name": "Sangam JayaVardhan Reddy", "proficiency": {"Big Data": {1: 0.38, 2: 0.43, 3: 0.48, 4: 0.53, 5: 0.58, 6: 0.63}}},
    "SEEE038": {"name": "Sarveshwaran S", "proficiency": {"ML Basics": {1: 0.71, 2: 0.76, 3: 0.81, 4: 0.86, 5: 0.91, 6: 0.96}}},
    "SEEE039": {"name": "Shanjay Sundhar S", "proficiency": {"IoT": {1: 0.68, 2: 0.73, 3: 0.78, 4: 0.83, 5: 0.88, 6: 0.93}}},
    "SEEE040": {"name": "Shreeya Kannan", "proficiency": {"Fuzzy Logic": {1: 0.62, 2: 0.67, 3: 0.72, 4: 0.77, 5: 0.82, 6: 0.87}}},
    "SEEE041": {"name": "Shreya S", "proficiency": {"AI Ethics": {1: 0.61, 2: 0.66, 3: 0.71, 4: 0.76, 5: 0.81, 6: 0.86}}},
    "SEEE042": {"name": "Sri Kamal Krishank S", "proficiency": {"Data Science": {1: 0.46, 2: 0.51, 3: 0.56, 4: 0.61, 5: 0.66, 6: 0.71}}},
    "SEEE043": {"name": "Sri Ramana S", "proficiency": {"Quantum Computing": {1: 0.37, 2: 0.42, 3: 0.47, 4: 0.52, 5: 0.57, 6: 0.62}}},
    "SEEE044": {"name": "Subashree S", "proficiency": {"Reinforcement Learning": {1: 0.48, 2: 0.53, 3: 0.58, 4: 0.63, 5: 0.68, 6: 0.73}}},
    "SEEE045": {"name": "Subhiksha N", "proficiency": {"Bayesian Learning": {1: 0.38, 2: 0.43, 3: 0.48, 4: 0.53, 5: 0.58, 6: 0.63}}},
    "SEEE046": {"name": "Sukhi Sudharshan S", "proficiency": {"Robotics": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE047": {"name": "Thanuja", "proficiency": {"Computer Vision": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}},
    "SEEE048": {"name": "Vasantha Kumar A", "proficiency": {"Quantum Computing": {1: 0.5, 2: 0.55, 3: 0.6, 4: 0.65, 5: 0.7, 6: 0.75}}},
    "SEEE049": {"name": "Velmugilan S", "proficiency": {"Embedded Systems": {1: 0.52, 2: 0.57, 3: 0.62, 4: 0.67, 5: 0.72, 6: 0.77}}},
    "SEEE050": {"name": "Velmurugan S", "proficiency": {"Fuzzy Logic": {1: 0.5, 2: 0.55, 3: 0.6, 4: 0.65, 5: 0.7, 6: 0.75}}},
    "SEEE051": {"name": "Viamrsh P S", "proficiency": {"Generative AI": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE052": {"name": "Vilohitan R", "proficiency": {"Reinforcement Learning": {1: 0.44, 2: 0.49, 3: 0.54, 4: 0.59, 5: 0.64, 6: 0.69}}},
    "SEEE053": {"name": "Vundela Guru Venkata Ajay Kumar Reddy", "proficiency": {"Big Data": {1: 0.72, 2: 0.77, 3: 0.82, 4: 0.87, 5: 0.92, 6: 0.97}}},
    "SEEE054": {"name": "Yadlapalli Madhuri", "proficiency": {"Robotics": {1: 0.44, 2: 0.49, 3: 0.54, 4: 0.59, 5: 0.64, 6: 0.69}}},
    "SEEE055": {"name": "Yohan K", "proficiency": {"Data Science": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE056": {"name": "Vinupriya A", "proficiency": {"Computer Vision": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE057": {"name": "Aparajita S", "proficiency": {"Computer Vision": {1: 0.73, 2: 0.78, 3: 0.83, 4: 0.88, 5: 0.93, 6: 0.98}}},
    "SEEE058": {"name": "Sriram L", "proficiency": {"NLP": {1: 0.63, 2: 0.68, 3: 0.73, 4: 0.78, 5: 0.83, 6: 0.88}}},
    "SEEE059": {"name": "Nandhini S", "proficiency": {"Deep Learning": {1: 0.55, 2: 0.6, 3: 0.65, 4: 0.7, 5: 0.75, 6: 0.8}}},
    "SEEE060": {"name": "Sreenath G", "proficiency": {"IoT": {1: 0.52, 2: 0.57, 3: 0.62, 4: 0.67, 5: 0.72, 6: 0.77}}},
    "SEEE061": {"name": "Swaran V", "proficiency": {"AI Ethics": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}},
    "SEEE061": {"name": "Swaran V", "proficiency": {"AI Ethics": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}}
}

# ---------------- SAMPLE QUESTIONS DATA ----------------
questions = [
    {"text": "What is the primary goal of supervised learning?", "options": [
        "To cluster data points without labels",
        "To learn patterns from labeled data to make predictions",
        "To reduce the dimensionality of data",
        "To generate new data samples"
    ], "correct": "To learn patterns from labeled data to make predictions"},
    {"text": "Which algorithm is used for classification?", "options": [
        "K-Means", "Decision Tree", "PCA", "Apriori"
    ], "correct": "Decision Tree"},
    {"text": "What does CNN stand for?", "options": [
        "Convolutional Neural Network", "Central Neural Network",
        "Computational Node Network", "Convolutional Node Network"
    ], "correct": "Convolutional Neural Network"},
    {"text": "Which technique helps reduce overfitting?", "options": [
        "Regularization", "Dropout", "Both 1 & 2", "None"
    ], "correct": "Both 1 & 2"}
]

# ---------------- SESSION STATE INIT ----------------
if "quiz_active" not in st.session_state:
    st.session_state.quiz_active = False
if "question_count" not in st.session_state:
    st.session_state.question_count = 0
if "total_questions" not in st.session_state:
    st.session_state.total_questions = 5
if "current_question" not in st.session_state:
    st.session_state.current_question = None
if "live_omr_questions" not in st.session_state:
    st.session_state.live_omr_questions = None
    
# Additional session states for enhanced OMR test
if "student_answers" not in st.session_state:
    st.session_state.student_answers = []
if "current_question_index" not in st.session_state:
    st.session_state.current_question_index = 0
if "warnings_count" not in st.session_state:
    st.session_state.warnings_count = 0
if "test_in_progress" not in st.session_state:
    st.session_state.test_in_progress = False
if "test_completed" not in st.session_state:
    st.session_state.test_completed = False
if "last_detection_time" not in st.session_state:
    st.session_state.last_detection_time = time.time()
if "looking_away_start" not in st.session_state:
    st.session_state.looking_away_start = None
if "audio_detection_queue" not in st.session_state:
    st.session_state.audio_detection_queue = queue.Queue()
# New states for enhanced proctoring
if "latest_warning" not in st.session_state:
    st.session_state.latest_warning = None
if "latest_warning_time" not in st.session_state:
    st.session_state.latest_warning_time = None
if "latest_warning_type" not in st.session_state:
    st.session_state.latest_warning_type = None
if "test_terminated" not in st.session_state:
    st.session_state.test_terminated = False
if "termination_reason" not in st.session_state:
    st.session_state.termination_reason = None
# Initialize AI Tutor
setup_ai_tutor()

# Initialize database
init_db()

# ---------------- STREAMLIT UI LAYOUT ----------------
st.title("🚀 Generative AI-Based Students Assessment System")
st.sidebar.header("Navigation")

# Main categories
main_category = st.sidebar.selectbox(
    "Main Menu",
    ["📊 Dashboard", "👨‍🏫 AI Tutor", "AI Exam", "AI Student Info", "AI Search", "AI Assessment"]
)

# Initialize page variable
page = None

if main_category == "📊 Dashboard":
    # Dashboard is directly accessible
    page = "📊 Dashboard"
    
elif main_category == "👨‍🏫 AI Tutor":
    # AI Tutor is directly accessible - just display the AI Tutor page
    display_ai_tutor_page(students_data)
    # Set page to None to prevent other pages from displaying
    page = None
    
elif main_category == "AI Exam":
    # Exam submenu
    exam = st.sidebar.selectbox(
        "Select Exam Type",
        [
           "📝 Take Quiz", 
           "📚 Quiz History",
           "📝 Live OMR Test",
        ]
    )
    page = exam
    
elif main_category == "AI Student Info":
    # Student info submenu
    student_info = st.sidebar.selectbox(
        "Select Student Tool",
        [
            "📄 AI-Based LOR Generator",
            "📈 ML Performance Tracking & Prediction"
        ]
    )
    page = student_info
    
elif main_category == "AI Search":
    # Search submenu
    search = st.sidebar.selectbox(
        "Select Search Tool",
        [
            "📖 AI-Powered Storytelling",
            "🔍 AI Peer Assessment",
            "🧠 AI-Powered Hints",
        ]
    )
    page = search
    
elif main_category == "AI Assessment":
    # Assessment submenu
    assessment = st.sidebar.selectbox(
        "Select Assessment Tool",
        [
            "🔍 Plagiarism/Reasoning Finder",
            "📂 Code Evaluation & Plagiarism Check",
            "📝 NLP Answer Evaluation",
            "📝 OMR MCQ Grading",
            "🔒 AI Proctoring & Integrity Checks"
        ]
    )
    page = assessment

            



# ---------------- PAGE: DASHBOARD ----------------
if page == "📊 Dashboard":
    st.header("📊 Class Performance Dashboard")
    
    # Select Student
    student_id = st.selectbox("Select a Student", list(students_data.keys()))
    student = students_data[student_id]
    st.subheader(f"Student: {student['name']}")
    
    # Proficiency Data
    prof_data = student["proficiency"]
    topics = list(prof_data.keys())
    bloom_levels = [1, 2, 3, 4, 5, 6]
    
    # Heatmap Data
    heatmap_data = pd.DataFrame(
        {topic: [prof_data[topic].get(level, 0.5) for level in bloom_levels] for topic in topics},
        index=[f"Level {lvl}" for lvl in bloom_levels]
    )
    
    st.subheader("📌 Bloom's Taxonomy Heatmap")
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.heatmap(heatmap_data, annot=True, cmap="coolwarm", linewidths=0.5, ax=ax)
    st.pyplot(fig)
    
    st.subheader("📈 Student Progress Over Time")
    dates = [datetime.now().replace(day=i) for i in range(1, 11)]
    scores = [random.uniform(0.4, 0.9) for _ in range(10)]
    plt.figure(figsize=(8, 5))
    plt.plot(dates, scores, marker='o', linestyle='-')
    plt.xticks(rotation=45)
    plt.xlabel("Date")
    plt.ylabel("Proficiency Score")
    plt.title("Student Progress")
    st.pyplot(plt)
# ---------------- PAGE: AI TUTOR ----------------
elif page == "👨‍🏫 AI Tutor":
    display_ai_tutor_page(students_data)
# ---------------- PAGE: AI-POWERED STORYTELLING ----------------
elif page == "📖 AI-Powered Storytelling":
    st.header("📖 AI-Powered Conversational Storytelling")
    user_topic = st.text_input("Enter a topic for the story (e.g., AI in Space, Lost Treasure, Cybersecurity):")
    if "story_conversation" not in st.session_state:
        st.session_state.story_conversation = []
    if st.button("Generate Story"):
        if user_topic.strip():
            prompt = (
                f"You are a friendly AI storyteller. Create a fun and engaging story based on the topic: '{user_topic}'. "
                f"Format the story like a chatbot conversation where an AI and a human character interact."
            )
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a storytelling AI."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.8,
                max_tokens=400
            )
            story_text = response.choices[0].message.content.strip()
            st.session_state.story_conversation.append({"role": "AI Storyteller", "content": story_text})
    for message in st.session_state.story_conversation:
        if message["role"] == "AI Storyteller":
            st.markdown(f"**🤖 AI Storyteller:** {message['content']}")
    user_reply = st.text_input("Continue the story... (Optional)")
    if st.button("Continue"):
        if user_reply.strip():
            follow_up_prompt = (
                f"Continue the story based on the user's response: '{user_reply}'. "
                f"Keep the conversational tone between the AI storyteller and the characters."
            )
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a storytelling AI."},
                    {"role": "user", "content": follow_up_prompt}
                ],
                temperature=0.8,
                max_tokens=300
            )
            follow_up_story = response.choices[0].message.content.strip()
            st.session_state.story_conversation.append({"role": "User", "content": user_reply})
            st.session_state.story_conversation.append({"role": "AI Storyteller", "content": follow_up_story})
            st.rerun()
# ---------------- PAGE: TAKE QUIZ ----------------
elif page == "📝 Take Quiz":
    st.header("📝 Adaptive Quiz")
    student_id = st.selectbox("Select your Student ID", list(students_data.keys()))
    st.write(f"Student Name: {students_data[student_id]['name']}")
    quiz_topic = st.selectbox(
        "Select a topic for your quiz:",
        ["Generative AI", "Machine Learning", "Reinforcement Learning", "Deep Learning", "Natural Language Processing", "Computer Vision"]
    )
    st.session_state.total_questions = st.slider("Number of Questions", min_value=1, max_value=10, value=5)
    if not st.session_state.quiz_active:
        if st.button("Start Quiz"):
            st.session_state.quiz_active = True
            st.session_state.question_count = 0
            st.session_state.score = 0
            prompt = (
                f"Create a multiple-choice question about {quiz_topic} with 4 options and indicate the correct answer. "
                f"Format your response as JSON with fields: 'text', 'options' (array of 4), and 'correct'."
            )
            with st.spinner("Generating question..."):
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are an AI that generates educational quiz questions. Return valid JSON only."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=250
                )
                try:
                    response_text = response.choices[0].message.content.strip()
                    if "```json" in response_text:
                        response_text = response_text.split("```json")[1].split("```")[0].strip()
                    elif "```" in response_text:
                        response_text = response_text.split("```")[1].split("```")[0].strip()
                    st.session_state.current_question = json.loads(response_text)
                    st.rerun()
                except Exception as e:
                    st.error(f"Error generating question: {str(e)}")
                    st.write("Response received:", response.choices[0].message.content)
                    st.session_state.quiz_active = False
    if st.session_state.quiz_active and st.session_state.question_count < st.session_state.total_questions:
        q = st.session_state.current_question
        st.write(f"**Q{st.session_state.question_count+1}:** {q['text']}")
        answer = st.radio("Choose the correct answer:", q["options"], index=None)
        if st.button("Submit Answer"):
            if answer:
                if answer == q["correct"]:
                    st.success("✅ Correct Answer!")
                    st.session_state.score += 1
                else:
                    st.error("❌ Incorrect Answer!")
                    st.info(f"The correct answer was: {q['correct']}")
                st.session_state.question_count += 1
                if st.session_state.question_count < st.session_state.total_questions:
                    prompt = (
                        f"Create a multiple-choice question about {quiz_topic} with 4 options and indicate the correct answer. "
                        f"Format as JSON: 'text', 'options', 'correct'."
                    )
                    with st.spinner("Generating next question..."):
                        response = openai.ChatCompletion.create(
                            model="gpt-4",
                            messages=[
                                {"role": "system", "content": "You are an AI that generates educational quiz questions. Return valid JSON only."},
                                {"role": "user", "content": prompt}
                            ],
                            temperature=0.7,
                            max_tokens=250
                        )
                        try:
                            response_text = response.choices[0].message.content.strip()
                            if "```json" in response_text:
                                response_text = response_text.split("```json")[1].split("```")[0].strip()
                            elif "```" in response_text:
                                response_text = response_text.split("```")[1].split("```")[0].strip()
                            st.session_state.current_question = json.loads(response_text)
                        except Exception as e:
                            st.error(f"Error generating question: {str(e)}")
                            st.write("Response received:", response.choices[0].message.content)
                            st.session_state.current_question = {
                                "text": "What is a primary advantage of using neural networks?",
                                "options": [
                                    "They always require less data than other models",
                                    "They can automatically learn complex patterns from data",
                                    "They never overfit",
                                    "They always train quickly"
                                ],
                                "correct": "They can automatically learn complex patterns from data"
                            }
                st.rerun()
    elif st.session_state.quiz_active and st.session_state.question_count >= st.session_state.total_questions:
        st.success(f"Quiz Complete! Your Final Score: {st.session_state.score}/{st.session_state.total_questions}")
        save_quiz_result(
            student_id=student_id,
            student_name=students_data[student_id]["name"],
            topic=quiz_topic,
            score=st.session_state.score,
            total_questions=st.session_state.total_questions
        )
        score_percentage = (st.session_state.score / st.session_state.total_questions) * 100
        if score_percentage >= 80:
            st.balloons()
            st.write("🌟 Excellent! You have a strong understanding of this topic.")
        elif score_percentage >= 60:
            st.write("👍 Good job! You have a solid grasp of the basics.")
        else:
            st.write("📚 Keep learning! Consider reviewing this topic more thoroughly.")
        if st.button("Restart Quiz"):
            st.session_state.quiz_active = False
            st.rerun()

# ---------------- PAGE: QUIZ HISTORY ----------------
elif page == "📚 Quiz History":
    st.header("📚 Quiz History Dashboard")
    col1, col2 = st.columns(2)
    with col1:
        view_option = st.radio("View", ["All Students", "Specific Student"])
    student_id = None
    if view_option == "Specific Student":
        with col2:
            student_id = st.selectbox("Select Student", list(students_data.keys()))
            st.write(f"Student: {students_data[student_id]['name']}")
    history_df = get_student_quiz_history(student_id)
    if not history_df.empty:
        history_df['percentage'] = (history_df['score'] / history_df['total_questions'] * 100).round(1)
        st.subheader("📋 Quiz Results")
        st.dataframe(history_df[['student_name', 'topic', 'score', 'total_questions', 'percentage', 'timestamp']])
        st.subheader("📊 Performance by Topic")
        topic_df = history_df.groupby('topic')[['percentage']].mean().reset_index()
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.barplot(x='topic', y='percentage', data=topic_df, ax=ax)
        plt.title("Average Score by Topic")
        plt.ylabel("Average Score (%)")
        plt.xlabel("Topic")
        plt.xticks(rotation=45)
        st.pyplot(fig)
        if student_id:
            st.subheader("📈 Performance Over Time")
            history_df['timestamp'] = pd.to_datetime(history_df['timestamp'])
            history_df = history_df.sort_values('timestamp')
            plt.figure(figsize=(10, 6))
            plt.plot(history_df['timestamp'], history_df['percentage'], marker='o', linestyle='-')
            plt.title(f"Performance Over Time - {students_data[student_id]['name']}")
            plt.ylabel("Score (%)")
            plt.xlabel("Date")
            plt.grid(True, linestyle='--', alpha=0.7)
            plt.xticks(rotation=45)
            st.pyplot(plt)
    else:
        st.info("No quiz history available. Take a quiz to see results here.")

# ---------------- PAGE: AI-POWERED HINTS ----------------
elif page == "🧠 AI-Powered Hints":
    st.header("🧠 Get AI-Powered Hints")
    question_text = st.text_area("Enter your question:")
    if st.button("Get Hint"):
        prompt = f"Provide a hint for the following question: {question_text}"
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an AI assistant providing hints."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
            max_tokens=100
        )
        hint = response.choices[0].message.content.strip()
        st.info(f"💡 Hint: {hint}")

# ---------------- PAGE: AI PEER ASSESSMENT ----------------
elif page == "🔍 AI Peer Assessment":
    st.header("🔍 AI-Powered Peer Assessment")
    student_submission = st.text_area("Paste student submission:")
    rubric = {"Concept Understanding": 10, "Implementation": 10, "Analysis": 10, "Clarity": 5, "Creativity": 5}
    if st.button("Generate Peer Review"):
        prompt = f"Assess the following submission using this rubric: {json.dumps(rubric)}\n\nSubmission: {student_submission}"
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an AI generating structured peer assessments."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=300
        )
        review = response.choices[0].message.content.strip()
        st.success("✅ AI Peer Review Generated:")
        st.write(review)

# ---------------- PAGE: PLAGIARISM / REASONING FINDER ----------------
elif page == "🔍 Plagiarism/Reasoning Finder":
    st.header("🔍 Plagiarism/Reasoning Finder")
    uploaded_file = st.file_uploader("Upload student's document (.docx or .pdf)", type=["docx", "pdf"])
    student_submission = ""
    if uploaded_file:
        file_type = uploaded_file.name.split(".")[-1].lower()
        if file_type == "docx":
            student_submission = extract_text_from_docx(uploaded_file)
        elif file_type == "pdf":
            student_submission = extract_text_from_pdf(uploaded_file)
        st.subheader("📄 Extracted Submission:")
        st.text_area("Extracted Text", student_submission, height=300)
    rubric = {"Concept Understanding": 10, "Implementation": 10, "Analysis": 10, "Clarity": 5, "Creativity": 5}
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📝 Generate AI Assessment"):
            if student_submission.strip():
                prompt = (
                    f"Assess the following submission based on this rubric: {json.dumps(rubric)}\n\n"
                    f"Submission:\n{student_submission}"
                )
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are an AI grading student assignments."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    max_tokens=500
                )
                review = response.choices[0].message.content.strip()
                st.success("✅ AI Feedback Generated:")
                st.write(review)
            else:
                st.warning("⚠️ Please upload a valid document.")
    with col2:
        if st.button("🔍 Check for Plagiarism"):
            if student_submission.strip():
                with st.spinner("Checking plagiarism..."):
                    plagiarism_prompt = (
                        f"Analyze the following document for plagiarism risk. "
                        f"Your response MUST include a clear plagiarism percentage in the first line with this exact format: "
                        f"'Plagiarism Risk: XX%' (where XX is a number between 0 and 100). "
                        f"Then provide your explanation below.\n\n"
                        f"Document:\n{student_submission}"
                    )
                    plagiarism_response = openai.ChatCompletion.create(
                        model="gpt-4",
                        messages=[
                            {
                                "role": "system",
                                "content": (
                                    "You are an AI that checks for plagiarism in academic work. "
                                    "Always provide a clear plagiarism percentage in your first line."
                                )
                            },
                            {"role": "user", "content": plagiarism_prompt}
                        ],
                        temperature=0.3,
                        max_tokens=400
                    )
                    plagiarism_result = plagiarism_response.choices[0].message.content.strip()
                    try:
                        if "Plagiarism Risk:" in plagiarism_result:
                            first_line = plagiarism_result.split('\n')[0]
                            percentage_text = first_line.split('Plagiarism Risk:')[1].strip()
                            if '%' in percentage_text:
                                percentage = float(percentage_text.replace('%', '').strip())
                                st.subheader("Plagiarism Risk Score:")
                                st.progress(percentage / 100)
                                if percentage < 30:
                                    st.success(f"📊 Plagiarism Risk: {percentage}% (Low Risk)")
                                elif percentage < 60:
                                    st.warning(f"📊 Plagiarism Risk: {percentage}% (Medium Risk)")
                                else:
                                    st.error(f"📊 Plagiarism Risk: {percentage}% (High Risk)")
                    except Exception as e:
                        st.error(f"Error parsing plagiarism percentage: {str(e)}")
                    st.subheader("📝 Plagiarism Analysis:")
                    st.write(plagiarism_result)
            else:
                st.warning("⚠️ Please upload a valid document.")

# ---------------- PAGE: CODE EVALUATION & PLAGIARISM CHECK ----------------
elif page == "📂 Code Evaluation & Plagiarism Check":
    st.header("📂 Code Evaluation & Plagiarism Check")
    uploaded_code = st.file_uploader("Upload a Python (.py) file", type=["py"])
    if uploaded_code is not None:
        code_content = uploaded_code.getvalue().decode("utf-8")
        st.subheader("📜 Uploaded Code:")
        st.code(code_content, language="python")
        syntax_error = check_python_syntax(code_content)
        if syntax_error:
            st.error(syntax_error)
        else:
            st.success("✅ No syntax errors found!")
        if st.button("📊 Evaluate Code & Check Plagiarism"):
            with st.spinner("Analyzing Code..."):
                prompt = (
                    "Analyze the following Python program. Provide:\n"
                    "1. Correctness Score\n2. Efficiency Score\n3. Readability Score\n"
                    "4. Plagiarism Risk Score\n5. Suggestions for Improvement\n\n"
                    f"Code:\n{code_content}"
                )
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are an AI that evaluates Python code."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    max_tokens=600
                )
                evaluation_result = response.choices[0].message.content.strip()
                st.success("✅ AI Evaluation Generated:")
                st.write(evaluation_result)
        if st.button("🔍 Check for Plagiarism"):
            with st.spinner("Checking plagiarism..."):
                plagiarism_prompt = (
                    f"Analyze the following Python code for plagiarism risk. "
                    f"Your response MUST include a clear plagiarism percentage in the first line with this exact format: "
                    f"'Plagiarism Risk: XX%' (where XX is 0-100). Then provide explanation below.\n\n"
                    f"Code:\n{code_content}"
                )
                plagiarism_response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {
                            "role": "system",
                            "content": (
                                "You are an AI that checks for plagiarism in Python code. "
                                "Always provide a clear plagiarism percentage in your first line."
                            )
                        },
                        {"role": "user", "content": plagiarism_prompt}
                    ],
                    temperature=0.3,
                    max_tokens=400
                )
                plagiarism_result = plagiarism_response.choices[0].message.content.strip()
                try:
                    if "Plagiarism Risk:" in plagiarism_result:
                        first_line = plagiarism_result.split('\n')[0]
                        percentage_text = first_line.split('Plagiarism Risk:')[1].strip()
                        if '%' in percentage_text:
                            percentage = float(percentage_text.replace('%', '').strip())
                            st.subheader("Plagiarism Risk Score:")
                            st.progress(percentage / 100)
                            if percentage < 30:
                                st.success(f"📊 Plagiarism Risk: {percentage}% (Low Risk)")
                            elif percentage < 60:
                                st.warning(f"📊 Plagiarism Risk: {percentage}% (Medium Risk)")
                            else:
                                st.error(f"📊 Plagiarism Risk: {percentage}% (High Risk)")
                except Exception as e:
                    st.error(f"Error parsing plagiarism percentage: {str(e)}")
                st.subheader("📝 Plagiarism Analysis:")
                st.write(plagiarism_result)

# ---------------- PAGE: AI-BASED LOR GENERATOR ----------------
elif page == "📄 AI-Based LOR Generator":
    st.header("📄 AI-Powered Letter of Recommendation (LOR) Generator")
    student_id = st.selectbox("Select Student ID", list(students_data.keys()))
    student = students_data[student_id]
    st.subheader(
        f"Student: {student['name']} | {student.get('department', 'B.Tech. Robotics & AI')}, "
        f"{student.get('year', 'Third Year')}"
    )
    st.write(f"**CGPA:** {student.get('cgpa', 'Not Available')}")
    lor_purpose = st.radio("Purpose of LOR", ["Higher Studies", "Internship", "Job Application"])
    skills = st.text_area("Enter key skills (comma-separated)", "Machine Learning, Research, Leadership")
    achievements = st.text_area("Enter achievements (comma-separated)", "Won AI Hackathon, Published IEEE Paper")
    if st.button("Generate LOR"):
        with st.spinner("Generating Letter of Recommendation..."):
            prompt = f"""
            Write a professional Letter of Recommendation (LOR) for a student. 
            The student details are:
            - Name: {student['name']}
            - Department: {student.get('department', 'B.Tech. Robotics & AI, Third Year')}
            - CGPA: {student.get('cgpa', 'Not Available')}
            - Purpose: {lor_purpose if lor_purpose else "Not Provided"}
            - Skills: {skills if skills else "Not Provided"}
            - Achievements: {achievements if achievements else "Not Provided"}

            The LOR should be formal, structured, and highlight the student's strengths, 
            academic achievements, and suitability for {lor_purpose}.
            """
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an AI that writes professional letters of recommendation."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.6,
                max_tokens=400
            )
            st.session_state.lor_text = response.choices[0].message.content.strip()
            st.success("✅ Letter of Recommendation Generated!")
            st.text_area("Generated LOR", st.session_state.lor_text, height=300)
    if "lor_text" in st.session_state:
        refine_prompt = st.text_area("🔄 Enter improvements to refine the LOR (optional)")
        if st.button("Refine LOR"):
            with st.spinner("Refining LOR..."):
                refine_request = f"Refine and improve the following LOR based on this feedback: {refine_prompt}\n\n{st.session_state.lor_text}"
                refined_response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You refine and improve letters of recommendation based on feedback."},
                        {"role": "user", "content": refine_request}
                    ],
                    temperature=0.6,
                    max_tokens=400
                )
                st.session_state.lor_text = refined_response.choices[0].message.content.strip()
                st.success("✅ LOR Refined!")
                st.text_area("Refined LOR", st.session_state.lor_text, height=300)
        doc = docx.Document()
        doc.add_paragraph(st.session_state.lor_text)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="📥 Download LOR (Word)",
            data=buffer,
            file_name="LOR.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ---------------- PAGE: NLP ANSWER EVALUATION ----------------
elif page == "📝 NLP Answer Evaluation":
    st.header("📝 NLP Answer Evaluation")
    question_prompt = st.text_area("Question/Prompt for the student:")
    student_answer = st.text_area("Student's Answer:")
    if st.button("Evaluate Answer"):
        if question_prompt.strip() and student_answer.strip():
            eval_prompt = (
                f"Evaluate the student's answer to the following question:\n\n"
                f"Question: {question_prompt}\n\n"
                f"Student Answer: {student_answer}\n\n"
                f"Provide a short evaluation of correctness, clarity, and completeness. Then provide a final grade (1-10)."
            )
            with st.spinner("Evaluating via GPT-4..."):
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are an AI specialized in NLP-based short-answer evaluation."},
                        {"role": "user", "content": eval_prompt}
                    ],
                    temperature=0.4,
                    max_tokens=250
                )
                evaluation = response.choices[0].message.content.strip()
            st.success("✅ Evaluation Complete:")
            st.write(evaluation)
        else:
            st.warning("Please provide both a question and a student answer.")

# ---------------- PAGE: OMR MCQ GRADING ----------------
elif page == "📝 OMR MCQ Grading":
    st.header("📝 Computer Vision-based OMR MCQ Grading")
    uploaded_omr = st.file_uploader("Upload a scanned OMR sheet (image)", type=["png", "jpg", "jpeg"])
    if uploaded_omr is not None:
        file_bytes = np.asarray(bytearray(uploaded_omr.read()), dtype=np.uint8)
        try:
            cv_image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
            st.image(cv_image, caption="Uploaded OMR Sheet", use_container_width=True)
            if st.button("Process OMR"):
                st.info("Processing OMR... (placeholder)")
                # TODO: Add actual OMR detection logic (e.g., circle detection, etc.)
                st.success("OMR processing complete. Placeholder result: Score = 8/10.")
        except Exception as e:
            st.error(f"Error reading image: {str(e)}")
    else:
        st.info("Upload an OMR image to begin.")


    # ---------------- PAGE: LIVE OMR TEST ----------------
elif page == "📝 Live OMR Test":
    # Use the enhanced OMR test function
    display_live_omr_test()

# ---------------- PAGE: ML PERFORMANCE TRACKING & PREDICTION ----------------
elif page == "📈 ML Performance Tracking & Prediction":
    st.header("📈 ML Performance Tracking & Prediction")
    st.write("This page illustrates how to use ML (Random Forest/XGBoost) to predict future performance.")
    st.markdown(
        """
        **Steps (conceptual):**
        1. Collect student performance data over time (scores, quiz results, etc.).
        2. Prepare features (e.g., recent quiz scores, attendance, etc.).
        3. Train a model (RandomForest/XGBoost) to predict future exam performance.
        4. Provide early interventions if the predicted performance is below a threshold.
        """
    )
    if st.button("Train Sample Model"):
        st.info("Training placeholder RandomForest model on synthetic data...")
        try:
            from sklearn.datasets import make_classification
            from sklearn.model_selection import train_test_split
            from sklearn.metrics import accuracy_score
            X, y = make_classification(n_samples=100, n_features=5, n_informative=3, random_state=42)
            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
            rf = RandomForestClassifier(n_estimators=10, random_state=42)
            rf.fit(X_train, y_train)
            preds = rf.predict(X_test)
            acc = accuracy_score(y_test, preds)
            st.success(f"Sample RandomForest Accuracy: {acc*100:.2f}%")
        except Exception as e:
            st.error("Scikit-learn not installed or an error occurred:")
            st.error(str(e))
    st.write("Similarly, you could use XGBoost for performance prediction:")
    if st.button("Train XGBoost Model"):
        try:
            from sklearn.datasets import make_regression
            from sklearn.model_selection import train_test_split
            from sklearn.metrics import r2_score
            X, y = make_regression(n_samples=100, n_features=5, n_informative=3, random_state=42)
            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
            model = xgb.XGBRegressor(n_estimators=10, use_label_encoder=False)
            model.fit(X_train, y_train, eval_set=[(X_test, y_test)], eval_metric='logloss')
            preds = model.predict(X_test)
            score = r2_score(y_test, preds)
            st.success(f"Sample XGBoost R2 Score: {score:.2f}")
        except Exception as e:
            st.error("XGBoost not installed or an error occurred:")
            st.error(str(e))
    st.info("Use the trained model to predict future performance and intervene early if needed.")

# ---------------- PAGE: AI PROCTORING & INTEGRITY CHECKS ----------------
elif page == "🔒 AI Proctoring & Integrity Checks":
    st.header("🔒 AI-Powered Proctoring & Integrity Checks")
    st.write("Below is a **conceptual** placeholder for real-time proctoring features.")
    st.markdown(
        """
        **Potential Approaches:**
        - **Facial Recognition**: Confirm the student's identity matches the ID photo.
        - **Eye-Tracking**: Check if eyes deviate from screen for suspicious durations.
        - **Behavior Analysis**: Detect multiple faces, unusual movements, etc.
        """
    )
    face_image = st.file_uploader("Upload a photo for face verification", type=["png", "jpg", "jpeg"])
    if face_image is not None:
        try:
            file_bytes = np.asarray(bytearray(face_image.read()), dtype=np.uint8)
            cv_image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
            if cv_image is not None:
                st.image(cv_image, caption="Uploaded Face Image", use_container_width=True)
                if st.button("Run Face Recognition"):
                    st.info("Performing face recognition... (placeholder)")
                    st.success("Identity verified with high confidence. (Placeholder result)")
            else:
                st.error("Error: Could not decode the image. Please try a different file.")
        except Exception as e:
            st.error(f"Error processing image: {str(e)}")
    st.write("For advanced proctoring (eye tracking, behavior analysis), you'd typically need a live camera feed.")
    st.write("Below is just a placeholder demonstration.")
    if st.button("Start Behavior Analysis"):
        st.info("Analyzing user behavior from webcam feed... (placeholder)")
        st.success("No suspicious behavior detected. (Placeholder)")
    st.warning("Note: Real-time proctoring requires additional setup (webcam access, advanced CV).")
