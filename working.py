import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import random
import openai
import json
from datetime import datetime
import fitz  # PyMuPDF for PDFs
import docx  # For Word document handling
import ast  # For checking Python syntax
import io  # For handling uploaded files
import sqlite3
import os
from io import BytesIO
from dotenv import load_dotenv
load_dotenv()

# openai.api_key = os.getenv("OPENAI_API_KEY")
openai.api_key = st.secrets["OPENAI_API_KEY"]
client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])  # Use Streamlit Secrets API Key

# Initialize Database
def init_db():
    conn = sqlite3.connect('student_quiz_history.db')
    c = conn.cursor()
    c.execute('''
    CREATE TABLE IF NOT EXISTS quiz_results (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        student_id TEXT NOT NULL,
        student_name TEXT NOT NULL,
        topic TEXT NOT NULL,
        score INTEGER NOT NULL,
        total_questions INTEGER NOT NULL,
        timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    conn.commit()
    return conn

# Save quiz results to database
def save_quiz_result(student_id, student_name, topic, score, total_questions):
    conn = init_db()
    c = conn.cursor()
    c.execute('''
    INSERT INTO quiz_results (student_id, student_name, topic, score, total_questions, timestamp)
    VALUES (?, ?, ?, ?, ?, ?)
    ''', (student_id, student_name, topic, score, total_questions, datetime.now()))
    conn.commit()
    conn.close()

# Get quiz history for a student
def get_student_quiz_history(student_id=None):
    conn = init_db()
    c = conn.cursor()
    
    if student_id:
        c.execute('''
        SELECT student_id, student_name, topic, score, total_questions, timestamp
        FROM quiz_results
        WHERE student_id = ?
        ORDER BY timestamp DESC
        ''', (student_id,))
    else:
        c.execute('''
        SELECT student_id, student_name, topic, score, total_questions, timestamp
        FROM quiz_results
        ORDER BY timestamp DESC
        ''')
    
    results = c.fetchall()
    conn.close()
    
    if results:
        return pd.DataFrame(results, columns=['student_id', 'student_name', 'topic', 'score', 'total_questions', 'timestamp'])
    return pd.DataFrame()

def extract_text_from_docx(uploaded_file):
    """Extract text from a .docx file"""
    try:
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])  # Extract paragraphs as text
        return text
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def extract_text_from_pdf(uploaded_file):
    """Extract text from a .pdf file."""
    try:
        text = ""
        with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text() + "\n"  # Extract text from each page
        return text
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def check_python_syntax(code):
    """Check for syntax errors in uploaded Python code."""
    try:
        ast.parse(code)
        return None  # No syntax errors
    except SyntaxError as e:
        return f"Syntax Error: {e}"

# Sample Student Data
students_data = {
    "SEEE001": {"name": "Adhithya V", "proficiency": {"Fuzzy Logic": {1: 0.59, 2: 0.64, 3: 0.69, 4: 0.74, 5: 0.79, 6: 0.84}}},
    "SEEE002": {"name": "Akety Manjunath", "proficiency": {"Data Science": {1: 0.6, 2: 0.65, 3: 0.7, 4: 0.75, 5: 0.8, 6: 0.85}}},
    "SEEE003": {"name": "Aravind S", "proficiency": {"Bayesian Learning": {1: 0.37, 2: 0.42, 3: 0.47, 4: 0.52, 5: 0.57, 6: 0.62}}},
    "SEEE004": {"name": "Aswin S", "proficiency": {"NLP": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}},
    "SEEE005": {"name": "Avinash Kumar R", "proficiency": {"Data Science": {1: 0.69, 2: 0.74, 3: 0.79, 4: 0.84, 5: 0.89, 6: 0.94}}},
    "SEEE007": {"name": "Bhavya Sri B", "proficiency": {"IoT": {1: 0.59, 2: 0.64, 3: 0.69, 4: 0.74, 5: 0.79, 6: 0.84}}},
    "SEEE008": {"name": "Challagandla Anantha Pavan", "proficiency": {"Cybersecurity": {1: 0.7, 2: 0.75, 3: 0.8, 4: 0.85, 5: 0.9, 6: 0.95}}},
    "SEEE009": {"name": "Nagadarahas Kumar C S", "proficiency": {"Cybersecurity": {1: 0.6, 2: 0.65, 3: 0.7, 4: 0.75, 5: 0.8, 6: 0.85}}},
    "SEEE010": {"name": "Dodda Sri Pujitha", "proficiency": {"Edge AI": {1: 0.73, 2: 0.78, 3: 0.83, 4: 0.88, 5: 0.93, 6: 0.98}}},
    "SEEE011": {"name": "Dondapati Vallapala Yami", "proficiency": {"Deep Learning": {1: 0.58, 2: 0.63, 3: 0.68, 4: 0.73, 5: 0.78, 6: 0.83}}},
    "SEEE012": {"name": "Guduru Venkata Sai Karthikeya", "proficiency": {"Generative AI": {1: 0.42, 2: 0.47, 3: 0.52, 4: 0.57, 5: 0.62, 6: 0.67}}},
    "SEEE013": {"name": "Hamsitha P", "proficiency": {"AI Ethics": {1: 0.39, 2: 0.44, 3: 0.49, 4: 0.54, 5: 0.59, 6: 0.64}}},
    "SEEE014": {"name": "Harini M", "proficiency": {"Generative AI": {1: 0.5, 2: 0.55, 3: 0.6, 4: 0.65, 5: 0.7, 6: 0.75}}},
    "SEEE015": {"name": "Harrish B", "proficiency": {"Robotics": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE016": {"name": "Jivan Prasath S", "proficiency": {"Swarm Intelligence": {1: 0.51, 2: 0.56, 3: 0.61, 4: 0.66, 5: 0.71, 6: 0.76}}},
    "SEEE017": {"name": "Jonnalagadda Susrith", "proficiency": {"Quantum Computing": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE018": {"name": "Kamaleshwar M", "proficiency": {"Data Science": {1: 0.72, 2: 0.77, 3: 0.82, 4: 0.87, 5: 0.92, 6: 0.97}}},
    "SEEE019": {"name": "Karthik Periyakarupphan M", "proficiency": {"NLP": {1: 0.36, 2: 0.41, 3: 0.46, 4: 0.51, 5: 0.56, 6: 0.61}}},
    "SEEE020": {"name": "Kathirazagan V", "proficiency": {"Bayesian Learning": {1: 0.44, 2: 0.49, 3: 0.54, 4: 0.59, 5: 0.64, 6: 0.69}}},
    "SEEE021": {"name": "Kishore R", "proficiency": {"Robotics": {1: 0.72, 2: 0.77, 3: 0.82, 4: 0.87, 5: 0.92, 6: 0.97}}},
    "SEEE022": {"name": "Krishnakumar Aditi J", "proficiency": {"Computer Vision": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE023": {"name": "Logavarshini K", "proficiency": {"Optimization": {1: 0.7, 2: 0.75, 3: 0.8, 4: 0.85, 5: 0.9, 6: 0.95}}},
    "SEEE024": {"name": "Malapalli Charitha Reddy", "proficiency": {"Deep Learning": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}},
    "SEEE025": {"name": "Manas M Nair", "proficiency": {"Bayesian Learning": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE026": {"name": "Mervath M", "proficiency": {"Robotics": {1: 0.55, 2: 0.6, 3: 0.65, 4: 0.7, 5: 0.75, 6: 0.8}}},
    "SEEE027": {"name": "Nidhish Kumar K", "proficiency": {"Bayesian Learning": {1: 0.37, 2: 0.42, 3: 0.47, 4: 0.52, 5: 0.57, 6: 0.62}}},
    "SEEE028": {"name": "Palapati Jahnavi", "proficiency": {"Reinforcement Learning": {1: 0.65, 2: 0.7, 3: 0.75, 4: 0.8, 5: 0.85, 6: 0.9}}},
    "SEEE029": {"name": "Pola Kaarthi", "proficiency": {"Fuzzy Logic": {1: 0.74, 2: 0.79, 3: 0.84, 4: 0.89, 5: 0.94, 6: 0.99}}},
    "SEEE030": {"name": "Pullela Vaishnavi", "proficiency": {"Data Science": {1: 0.39, 2: 0.44, 3: 0.49, 4: 0.54, 5: 0.59, 6: 0.64}}},
    "SEEE031": {"name": "Raghul T", "proficiency": {"Generative AI": {1: 0.43, 2: 0.48, 3: 0.53, 4: 0.58, 5: 0.63, 6: 0.68}}},
    "SEEE032": {"name": "Rajaprabu K", "proficiency": {"Big Data": {1: 0.41, 2: 0.46, 3: 0.51, 4: 0.56, 5: 0.61, 6: 0.66}}},
    "SEEE033": {"name": "Rishi A", "proficiency": {"NLP": {1: 0.41, 2: 0.46, 3: 0.51, 4: 0.56, 5: 0.61, 6: 0.66}}},
    "SEEE034": {"name": "Rithick Reddy S", "proficiency": {"Swarm Intelligence": {1: 0.48, 2: 0.53, 3: 0.58, 4: 0.63, 5: 0.68, 6: 0.73}}},
    "SEEE035": {"name": "Tera Sai Tejeshwar Reddy", "proficiency": {"Swarm Intelligence": {1: 0.59, 2: 0.64, 3: 0.69, 4: 0.74, 5: 0.79, 6: 0.84}}},
    "SEEE036": {"name": "Saikirthiga R", "proficiency": {"Data Science": {1: 0.55, 2: 0.6, 3: 0.65, 4: 0.7, 5: 0.75, 6: 0.8}}},
    "SEEE037": {"name": "Sangam JayaVardhan Reddy", "proficiency": {"Big Data": {1: 0.38, 2: 0.43, 3: 0.48, 4: 0.53, 5: 0.58, 6: 0.63}}},
    "SEEE038": {"name": "Sarveshwaran S", "proficiency": {"ML Basics": {1: 0.71, 2: 0.76, 3: 0.81, 4: 0.86, 5: 0.91, 6: 0.96}}},
    "SEEE039": {"name": "Shanjay Sundhar S", "proficiency": {"IoT": {1: 0.68, 2: 0.73, 3: 0.78, 4: 0.83, 5: 0.88, 6: 0.93}}},
    "SEEE040": {"name": "Shreeya Kannan", "proficiency": {"Fuzzy Logic": {1: 0.62, 2: 0.67, 3: 0.72, 4: 0.77, 5: 0.82, 6: 0.87}}},
    "SEEE041": {"name": "Shreya S", "proficiency": {"AI Ethics": {1: 0.61, 2: 0.66, 3: 0.71, 4: 0.76, 5: 0.81, 6: 0.86}}},
    "SEEE042": {"name": "Sri Kamal Krishank S", "proficiency": {"Data Science": {1: 0.46, 2: 0.51, 3: 0.56, 4: 0.61, 5: 0.66, 6: 0.71}}},
    "SEEE043": {"name": "Sri Ramana S", "proficiency": {"Quantum Computing": {1: 0.37, 2: 0.42, 3: 0.47, 4: 0.52, 5: 0.57, 6: 0.62}}},
    "SEEE044": {"name": "Subashree S", "proficiency": {"Reinforcement Learning": {1: 0.48, 2: 0.53, 3: 0.58, 4: 0.63, 5: 0.68, 6: 0.73}}},
    "SEEE045": {"name": "Subhiksha N", "proficiency": {"Bayesian Learning": {1: 0.38, 2: 0.43, 3: 0.48, 4: 0.53, 5: 0.58, 6: 0.63}}},
    "SEEE046": {"name": "Sukhi Sudharshan S", "proficiency": {"Robotics": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE047": {"name": "Thanuja", "proficiency": {"Computer Vision": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}},
    "SEEE048": {"name": "Vasantha Kumar A", "proficiency": {"Quantum Computing": {1: 0.5, 2: 0.55, 3: 0.6, 4: 0.65, 5: 0.7, 6: 0.75}}},
    "SEEE049": {"name": "Velmugilan S", "proficiency": {"Embedded Systems": {1: 0.52, 2: 0.57, 3: 0.62, 4: 0.67, 5: 0.72, 6: 0.77}}},
    "SEEE050": {"name": "Velmurugan S", "proficiency": {"Fuzzy Logic": {1: 0.5, 2: 0.55, 3: 0.6, 4: 0.65, 5: 0.7, 6: 0.75}}},
    "SEEE051": {"name": "Viamrsh P S", "proficiency": {"Generative AI": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE052": {"name": "Vilohitan R", "proficiency": {"Reinforcement Learning": {1: 0.44, 2: 0.49, 3: 0.54, 4: 0.59, 5: 0.64, 6: 0.69}}},
    "SEEE053": {"name": "Vundela Guru Venkata Ajay Kumar Reddy", "proficiency": {"Big Data": {1: 0.72, 2: 0.77, 3: 0.82, 4: 0.87, 5: 0.92, 6: 0.97}}},
    "SEEE054": {"name": "Yadlapalli Madhuri", "proficiency": {"Robotics": {1: 0.44, 2: 0.49, 3: 0.54, 4: 0.59, 5: 0.64, 6: 0.69}}},
    "SEEE055": {"name": "Yohan K", "proficiency": {"Data Science": {1: 0.64, 2: 0.69, 3: 0.74, 4: 0.79, 5: 0.84, 6: 0.89}}},
    "SEEE056": {"name": "Vinupriya A", "proficiency": {"Computer Vision": {1: 0.45, 2: 0.5, 3: 0.55, 4: 0.6, 5: 0.65, 6: 0.7}}},
    "SEEE057": {"name": "Aparajita S", "proficiency": {"Computer Vision": {1: 0.73, 2: 0.78, 3: 0.83, 4: 0.88, 5: 0.93, 6: 0.98}}},
    "SEEE058": {"name": "Sriram L", "proficiency": {"NLP": {1: 0.63, 2: 0.68, 3: 0.73, 4: 0.78, 5: 0.83, 6: 0.88}}},
    "SEEE059": {"name": "Nandhini S", "proficiency": {"Deep Learning": {1: 0.55, 2: 0.6, 3: 0.65, 4: 0.7, 5: 0.75, 6: 0.8}}},
    "SEEE060": {"name": "Sreenath G", "proficiency": {"IoT": {1: 0.52, 2: 0.57, 3: 0.62, 4: 0.67, 5: 0.72, 6: 0.77}}},
    "SEEE061": {"name": "Swaran V", "proficiency": {"AI Ethics": {1: 0.54, 2: 0.59, 3: 0.64, 4: 0.69, 5: 0.74, 6: 0.79}}}

}



# Sample Questions Data
questions = [
    {"text": "What is the primary goal of supervised learning?", "options": [
        "To cluster data points without labels",
        "To learn patterns from labeled data to make predictions",
        "To reduce the dimensionality of data",
        "To generate new data samples"
    ], "correct": "To learn patterns from labeled data to make predictions"},
    {"text": "Which algorithm is used for classification?", "options": [
        "K-Means", "Decision Tree", "PCA", "Apriori"
    ], "correct": "Decision Tree"},
    {"text": "What does CNN stand for?", "options": [
        "Convolutional Neural Network", "Central Neural Network",
        "Computational Node Network", "Convolutional Node Network"
    ], "correct": "Convolutional Neural Network"},
    {"text": "Which technique helps reduce overfitting?", "options": [
        "Regularization", "Dropout", "Both 1 & 2", "None"
    ], "correct": "Both 1 & 2"}
]

# Initialize session states for quiz tracking
if "quiz_active" not in st.session_state:
    st.session_state.quiz_active = False
if "question_count" not in st.session_state:
    st.session_state.question_count = 0
if "total_questions" not in st.session_state:
    st.session_state.total_questions = 5
if "current_question" not in st.session_state:
    st.session_state.current_question = None
if "score" not in st.session_state:
    st.session_state.score = 0

# Initialize database
init_db()

# Streamlit UI
st.title("🚀 Generative AI-Based Students Assessment System")
st.sidebar.header("Navigation")
page = st.sidebar.radio("Go to", ["📊 Dashboard", "📝 Take Quiz", "📚 Quiz History", "📖 AI-Powered Storytelling", "🧠 AI-Powered Hints", "🔍 AI Peer Assessment", "🔍 Plagiarism/Reasoning Finder", "📂 Code Evaluation & Plagiarism Check", "📄 AI-Based LOR Generator"
])
### 📊 CLASS PERFORMANCE DASHBOARD ###
if page == "📊 Dashboard":
    st.header("📊 Class Performance Dashboard")
    
    # Select Student
    student_id = st.selectbox("Select a Student", list(students_data.keys()))
    student = students_data[student_id]
    st.subheader(f"Student: {student['name']}")
    
    # Proficiency Data
    prof_data = student["proficiency"]
    topics = list(prof_data.keys())
    bloom_levels = [1, 2, 3, 4, 5, 6]
    
    # Heatmap Data
    heatmap_data = pd.DataFrame({topic: [prof_data[topic].get(level, 0.5) for level in bloom_levels] for topic in topics},
                                 index=[f"Level {lvl}" for lvl in bloom_levels])
    
    st.subheader("📌 Bloom's Taxonomy Heatmap")
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.heatmap(heatmap_data, annot=True, cmap="coolwarm", linewidths=0.5, ax=ax)
    st.pyplot(fig)
    
    st.subheader("📈 Student Progress Over Time")
    dates = [datetime.now().replace(day=i) for i in range(1, 11)]
    scores = [random.uniform(0.4, 0.9) for _ in range(10)]
    plt.figure(figsize=(8, 5))
    plt.plot(dates, scores, marker='o', linestyle='-')
    plt.xticks(rotation=45)
    plt.xlabel("Date")
    plt.ylabel("Proficiency Score")
    plt.title("Student Progress")
    st.pyplot(plt)
if page == "📖 AI-Powered Storytelling":
    st.header("📖 AI-Powered Conversational Storytelling")

    # User input for story topic
    user_topic = st.text_input("Enter a topic for the story (e.g., AI in Space, Lost Treasure, Cybersecurity):")

    if "story_conversation" not in st.session_state:
        st.session_state.story_conversation = []

    if st.button("Generate Story"):
        if user_topic.strip():
            # OpenAI API call to generate the first part of the story
            prompt = (
                f"You are a friendly AI storyteller. Create a fun and engaging story based on the topic: '{user_topic}'. "
                f"Format the story like a chatbot conversation where an AI and a human character interact."
            )

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "system", "content": "You are a storytelling AI."},
                          {"role": "user", "content": prompt}],
                temperature=0.8,
                max_tokens=400
            )

            story_text = response.choices[0].message.content.strip()

            # Add story to session history
            st.session_state.story_conversation.append({"role": "AI Storyteller", "content": story_text})

    # Display the chatbot-style story conversation
    for message in st.session_state.story_conversation:
        if message["role"] == "AI Storyteller":
            st.markdown(f"**🤖 AI Storyteller:** {message['content']}")

    # Allow users to add their own input to progress the story
    user_reply = st.text_input("Continue the story... (Optional)")

    if st.button("Continue"):
        if user_reply.strip():
            follow_up_prompt = (
                f"Continue the story based on the user's response: '{user_reply}'. "
                f"Keep the conversational tone between the AI storyteller and the characters."
            )

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "system", "content": "You are a storytelling AI."},
                          {"role": "user", "content": follow_up_prompt}],
                temperature=0.8,
                max_tokens=300
            )

            follow_up_story = response.choices[0].message.content.strip()

            # Append to conversation history
            st.session_state.story_conversation.append({"role": "User", "content": user_reply})
            st.session_state.story_conversation.append({"role": "AI Storyteller", "content": follow_up_story})

            # Refresh page
            st.rerun()

### 📝 DYNAMIC QUIZ GENERATION ###
elif page == "📝 Take Quiz":
    st.header("📝 Adaptive Quiz")
    
    # Student selection
    student_id = st.selectbox("Select your Student ID", list(students_data.keys()))
    st.write(f"Student Name: {students_data[student_id]['name']}")
    
    # Topic selection
    quiz_topic = st.selectbox(
        "Select a topic for your quiz:",
        ["Generative AI", "Machine Learning", "Reinforcement Learning", "Deep Learning", "Natural Language Processing", "Computer Vision"]
    )
    
    # Select number of questions
    st.session_state.total_questions = st.slider("Number of Questions", min_value=1, max_value=10, value=5)
    
    # Start Quiz
    if not st.session_state.quiz_active:
        if st.button("Start Quiz"):
            st.session_state.quiz_active = True
            st.session_state.question_count = 0
            st.session_state.score = 0
            
            # Generate the first question using OpenAI
            prompt = f"Create a multiple-choice question about {quiz_topic} with 4 options and indicate the correct answer. Format your response as JSON with these fields: 'text' (the question), 'options' (array of 4 choices), and 'correct' (the correct answer text)."
            
            with st.spinner("Generating question..."):
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "You are an AI that generates educational quiz questions. Return responses in valid JSON format only."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=250
                )
                
                # Parse the response
                try:
                    response_text = response.choices[0].message.content.strip()
                    # Clean up response if it contains extra markdown code blocks
                    if "```json" in response_text:
                        response_text = response_text.split("```json")[1].split("```")[0].strip()
                    elif "```" in response_text:
                        response_text = response_text.split("```")[1].split("```")[0].strip()
                    
                    st.session_state.current_question = json.loads(response_text)
                    st.rerun()  # Use st.rerun() instead of st.experimental_rerun()
                except Exception as e:
                    st.error(f"Error generating question: {str(e)}")
                    st.write("Response received:", response.choices[0].message.content)
                    st.session_state.quiz_active = False

    # Quiz Progress
    if st.session_state.quiz_active and st.session_state.question_count < st.session_state.total_questions:
        q = st.session_state.current_question
        st.write(f"**Q{st.session_state.question_count+1}:** {q['text']}")
        answer = st.radio("Choose the correct answer:", q["options"], index=None)
        
        if st.button("Submit Answer"):
            if answer:
                if answer == q["correct"]:
                    st.success("✅ Correct Answer!")
                    st.session_state.score += 1
                else:
                    st.error("❌ Incorrect Answer!")
                    st.info(f"The correct answer was: {q['correct']}")
                
                st.session_state.question_count += 1
                
                # Generate the next question if not done
                if st.session_state.question_count < st.session_state.total_questions:
                    prompt = f"Create a multiple-choice question about {quiz_topic} with 4 options and indicate the correct answer. Format your response as JSON with these fields: 'text' (the question), 'options' (array of 4 choices), and 'correct' (the correct answer text)."
                    
                    with st.spinner("Generating next question..."):
                        response = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {"role": "system", "content": "You are an AI that generates educational quiz questions. Return responses in valid JSON format only."},
                                {"role": "user", "content": prompt}
                            ],
                            temperature=0.7,
                            max_tokens=250
                        )
                        
                        # Parse the response
                        try:
                            response_text = response.choices[0].message.content.strip()
                            # Clean up response if it contains extra markdown code blocks
                            if "```json" in response_text:
                                response_text = response_text.split("```json")[1].split("```")[0].strip()
                            elif "```" in response_text:
                                response_text = response_text.split("```")[1].split("```")[0].strip()
                            
                            st.session_state.current_question = json.loads(response_text)
                        except Exception as e:
                            st.error(f"Error generating question: {str(e)}")
                            st.write("Response received:", response.choices[0].message.content)
                            # Provide a backup question in case of parsing error
                            st.session_state.current_question = {
                                "text": "What is a primary advantage of using neural networks?",
                                "options": [
                                    "They always require less data than other models",
                                    "They can automatically learn complex patterns from data",
                                    "They never overfit",
                                    "They always train quickly"
                                ],
                                "correct": "They can automatically learn complex patterns from data"
                            }
                
                st.rerun()  # Use st.rerun() instead of st.experimental_rerun()

    # Quiz completion
    elif st.session_state.quiz_active and st.session_state.question_count >= st.session_state.total_questions:
        st.success(f"Quiz Complete! Your Final Score: {st.session_state.score}/{st.session_state.total_questions}")
        
        # Save quiz results to database
        save_quiz_result(
            student_id=student_id,
            student_name=students_data[student_id]["name"],
            topic=quiz_topic,
            score=st.session_state.score,
            total_questions=st.session_state.total_questions
        )
        
        # Generate feedback based on score
        score_percentage = (st.session_state.score / st.session_state.total_questions) * 100
        if score_percentage >= 80:
            st.balloons()
            st.write("🌟 Excellent! You have a strong understanding of this topic.")
        elif score_percentage >= 60:
            st.write("👍 Good job! You have a solid grasp of the basics.")
        else:
            st.write("📚 Keep learning! Consider reviewing this topic more thoroughly.")
        
        if st.button("Restart Quiz"):
            st.session_state.quiz_active = False
            st.rerun()

### 📚 QUIZ HISTORY ###
elif page == "📚 Quiz History":
    st.header("📚 Quiz History Dashboard")
    
    # Filter options
    col1, col2 = st.columns(2)
    
    with col1:
        view_option = st.radio("View", ["All Students", "Specific Student"])
    
    student_id = None
    if view_option == "Specific Student":
        with col2:
            student_id = st.selectbox("Select Student", list(students_data.keys()))
            st.write(f"Student: {students_data[student_id]['name']}")
    
    # Fetch and display quiz history
    history_df = get_student_quiz_history(student_id)
    
    if not history_df.empty:
        # Add percentage column
        history_df['percentage'] = (history_df['score'] / history_df['total_questions'] * 100).round(1)
        
        # Display results table
        st.subheader("📋 Quiz Results")
        st.dataframe(history_df[['student_name', 'topic', 'score', 'total_questions', 'percentage', 'timestamp']])
        
        # Visualization - Performance by Topic
        st.subheader("📊 Performance by Topic")
        topic_df = history_df.groupby('topic')[['percentage']].mean().reset_index()
        
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.barplot(x='topic', y='percentage', data=topic_df, ax=ax)
        plt.title("Average Score by Topic")
        plt.ylabel("Average Score (%)")
        plt.xlabel("Topic")
        plt.xticks(rotation=45)
        st.pyplot(fig)
        
        # Visualization - Performance Over Time
        if student_id:
            st.subheader("📈 Performance Over Time")
            history_df['timestamp'] = pd.to_datetime(history_df['timestamp'])
            history_df = history_df.sort_values('timestamp')
            
            plt.figure(figsize=(10, 6))
            plt.plot(history_df['timestamp'], history_df['percentage'], marker='o', linestyle='-')
            plt.title(f"Performance Over Time - {students_data[student_id]['name']}")
            plt.ylabel("Score (%)")
            plt.xlabel("Date")
            plt.grid(True, linestyle='--', alpha=0.7)
            plt.xticks(rotation=45)
            st.pyplot(plt)
    else:
        st.info("No quiz history available. Take a quiz to see results here.")

### 🧠 AI POWERED HINTS ###
elif page == "🧠 AI-Powered Hints":
    st.header("🧠 Get AI-Powered Hints")
    question_text = st.text_area("Enter your question:")
    
    if st.button("Get Hint"):
        prompt = f"Provide a hint for the following question: {question_text}"
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": "You are an AI assistant providing hints."},
                      {"role": "user", "content": prompt}],
            temperature=0.4, max_tokens=100
        )
        hint = response.choices[0].message.content.strip()
        st.info(f"💡 Hint: {hint}")

### 🔍 AI PEER ASSESSMENT ###
elif page == "🔍 AI Peer Assessment":
    st.header("🔍 AI-Powered Peer Assessment")
    student_submission = st.text_area("Paste student submission:")
    rubric = {"Concept Understanding": 10, "Implementation": 10, "Analysis": 10, "Clarity": 5, "Creativity": 5}
    
    if st.button("Generate Peer Review"):
        prompt = f"Assess the following submission using this rubric: {json.dumps(rubric)}\n\nSubmission: {student_submission}"
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": "You are an AI generating structured peer assessments."},
                      {"role": "user", "content": prompt}],
            temperature=0.3, max_tokens=300
        )
        review = response.choices[0].message.content.strip()
        st.success("✅ AI Peer Review Generated:")
        st.write(review)

### 🔍 AI PEER ASSESSMENT WITH WORD DOCUMENT UPLOAD ###
### 🔍 AI PEER ASSESSMENT WITH WORD DOCUMENT UPLOAD ###
elif page == "🔍 Plagiarism/Reasoning Finder":
    st.header("🔍 Plagiarism/Reasoning Finder")
    uploaded_file = st.file_uploader("Upload student's document (.docx or .pdf)", type=["docx", "pdf"])

    student_submission = ""
    if uploaded_file:
        file_type = uploaded_file.name.split(".")[-1].lower()

        if file_type == "docx":
            student_submission = extract_text_from_docx(uploaded_file)
        elif file_type == "pdf":
            student_submission = extract_text_from_pdf(uploaded_file)

        st.subheader("📄 Extracted Submission:")
        st.text_area("Extracted Text", student_submission, height=300)

    rubric = {"Concept Understanding": 10, "Implementation": 10, "Analysis": 10, "Clarity": 5, "Creativity": 5}

    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📝 Generate AI Assessment"):
            if student_submission.strip():
                prompt = f"Assess the following submission based on this rubric: {json.dumps(rubric)}\n\nSubmission:\n{student_submission}"

                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "system", "content": "You are an AI grading student assignments."},
                              {"role": "user", "content": prompt}],
                    temperature=0.3, max_tokens=500
                )

                review = response.choices[0].message.content.strip()
                st.success("✅ AI Feedback Generated:")
                st.write(review)
            else:
                st.warning("⚠️ Please upload a valid document.")
    
    with col2:
        if st.button("🔍 Check for Plagiarism"):
            if student_submission.strip():
                with st.spinner("Checking plagiarism..."):
                    # More structured prompt to ensure we get a percentage
                    plagiarism_prompt = (
                        f"Analyze the following document for plagiarism risk. "
                        f"Your response MUST include a clear plagiarism percentage in the first line with this exact format: 'Plagiarism Risk: XX%' "
                        f"(where XX is a number between 0 and 100). Then provide your explanation and analysis below that.\n\n"
                        f"Document:\n{student_submission}"
                    )

                    plagiarism_response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "You are an AI that checks for plagiarism in academic work. Always provide a clear plagiarism percentage in your first line."},
                            {"role": "user", "content": plagiarism_prompt}
                        ],
                        temperature=0.3, max_tokens=400
                    )

                    plagiarism_result = plagiarism_response.choices[0].message.content.strip()
                    
                    # Extract the percentage for visualization
                    try:
                        if "Plagiarism Risk:" in plagiarism_result:
                            first_line = plagiarism_result.split('\n')[0]
                            percentage_text = first_line.split('Plagiarism Risk:')[1].strip()
                            if '%' in percentage_text:
                                percentage = float(percentage_text.replace('%', '').strip())
                                
                                # Create a progress bar to visualize the plagiarism percentage
                                st.subheader("Plagiarism Risk Score:")
                                st.progress(percentage/100)
                                
                                # Color-coded display based on risk level
                                if percentage < 30:
                                    st.success(f"📊 Plagiarism Risk: {percentage}% (Low Risk)")
                                elif percentage < 60:
                                    st.warning(f"📊 Plagiarism Risk: {percentage}% (Medium Risk)")
                                else:
                                    st.error(f"📊 Plagiarism Risk: {percentage}% (High Risk)")
                    except Exception as e:
                        st.error(f"Error parsing plagiarism percentage: {str(e)}")
                    
                    # Display full analysis
                    st.subheader("📝 Plagiarism Analysis:")
                    st.write(plagiarism_result)
                    
                    # Store the result in database if needed
                    if "current_student_id" in st.session_state and uploaded_file:
                        try:
                            conn = init_db()
                            c = conn.cursor()
                            # Create document_plagiarism_checks table if it doesn't exist
                            c.execute('''
                            CREATE TABLE IF NOT EXISTS document_plagiarism_checks (
                                id INTEGER PRIMARY KEY AUTOINCREMENT,
                                student_id TEXT,
                                filename TEXT,
                                plagiarism_score REAL,
                                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                            )
                            ''')
                            
                            # Try to extract the percentage for database storage
                            percentage = 0
                            try:
                                if "Plagiarism Risk:" in plagiarism_result:
                                    first_line = plagiarism_result.split('\n')[0]
                                    percentage_text = first_line.split('Plagiarism Risk:')[1].strip()
                                    if '%' in percentage_text:
                                        percentage = float(percentage_text.replace('%', '').strip())
                            except:
                                percentage = 0
                                
                            # Store the check result
                            c.execute('''
                            INSERT INTO document_plagiarism_checks (student_id, filename, plagiarism_score, timestamp)
                            VALUES (?, ?, ?, ?)
                            ''', (st.session_state.current_student_id, uploaded_file.name, percentage, datetime.now()))
                            
                            conn.commit()
                            conn.close()
                        except Exception as e:
                            st.error(f"Error saving plagiarism check to database: {str(e)}")
            else:
                st.warning("⚠️ Please upload a valid document.")
### 📂 CODE EVALUATION & PLAGIARISM CHECK ###
elif page == "📂 Code Evaluation & Plagiarism Check":
    st.header("📂 Code Evaluation & Plagiarism Check")

    uploaded_code = st.file_uploader("Upload a Python (.py) file", type=["py"])
    
    if uploaded_code is not None:
        # Read the uploaded file
        code_content = uploaded_code.getvalue().decode("utf-8")
        
        st.subheader("📜 Uploaded Code:")
        st.code(code_content, language="python")

        # Syntax Checking
        syntax_error = check_python_syntax(code_content)
        if syntax_error:
            st.error(syntax_error)
        else:
            st.success("✅ No syntax errors found!")

        # AI Code Evaluation
        if st.button("📊 Evaluate Code & Check Plagiarism"):
            with st.spinner("Analyzing Code..."):
                prompt = f"Analyze the following Python program. Provide:\n1. Correctness Score\n2. Efficiency Score\n3. Readability Score\n4. Plagiarism Risk Score\n5. Suggestions for Improvement\n\nCode:\n{code_content}"

                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "system", "content": "You are an AI that evaluates Python code for correctness, efficiency, readability, and plagiarism risk."},
                              {"role": "user", "content": prompt}],
                    temperature=0.3, max_tokens=600
                )

                evaluation_result = response.choices[0].message.content.strip()
                st.success("✅ AI Evaluation Generated:")
                st.write(evaluation_result)

        # Enhanced Plagiarism Score Analysis
        if st.button("🔍 Check for Plagiarism"):
            with st.spinner("Checking plagiarism..."):
                # More structured prompt to ensure we get a percentage
                plagiarism_prompt = (
                    f"Analyze the following Python code for plagiarism risk. "
                    f"Your response MUST include a clear plagiarism percentage in the first line with this exact format: 'Plagiarism Risk: XX%' "
                    f"(where XX is a number between 0 and 100). Then provide your explanation and analysis below that.\n\n"
                    f"Code:\n{code_content}"
                )

                plagiarism_response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "You are an AI that checks for plagiarism in Python code. Always provide a clear plagiarism percentage in your first line."},
                        {"role": "user", "content": plagiarism_prompt}
                    ],
                    temperature=0.3, max_tokens=400
                )

                plagiarism_result = plagiarism_response.choices[0].message.content.strip()
                
                # Extract the percentage for visualization
                try:
                    if "Plagiarism Risk:" in plagiarism_result:
                        first_line = plagiarism_result.split('\n')[0]
                        percentage_text = first_line.split('Plagiarism Risk:')[1].strip()
                        if '%' in percentage_text:
                            percentage = float(percentage_text.replace('%', '').strip())
                            
                            # Create a progress bar to visualize the plagiarism percentage
                            st.subheader("Plagiarism Risk Score:")
                            st.progress(percentage/100)
                            
                            # Color-coded display based on risk level
                            if percentage < 30:
                                st.success(f"📊 Plagiarism Risk: {percentage}% (Low Risk)")
                            elif percentage < 60:
                                st.warning(f"📊 Plagiarism Risk: {percentage}% (Medium Risk)")
                            else:
                                st.error(f"📊 Plagiarism Risk: {percentage}% (High Risk)")
                except Exception as e:
                    st.error(f"Error parsing plagiarism percentage: {str(e)}")
                
                # Display full analysis
                st.subheader("📝 Plagiarism Analysis:")
                st.write(plagiarism_result)
                
                # Store the result in database
                if "current_student_id" in st.session_state:
                    try:
                        conn = init_db()
                        c = conn.cursor()
                        # Create plagiarism_checks table if it doesn't exist
                        c.execute('''
                        CREATE TABLE IF NOT EXISTS plagiarism_checks (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            student_id TEXT,
                            filename TEXT,
                            plagiarism_score REAL,
                            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                        )
                        ''')
                        
                        # Try to extract the percentage for database storage
                        percentage = 0
                        try:
                            if "Plagiarism Risk:" in plagiarism_result:
                                first_line = plagiarism_result.split('\n')[0]
                                percentage_text = first_line.split('Plagiarism Risk:')[1].strip()
                                if '%' in percentage_text:
                                    percentage = float(percentage_text.replace('%', '').strip())
                        except:
                            percentage = 0
                            
                        # Store the check result
                        c.execute('''
                        INSERT INTO plagiarism_checks (student_id, filename, plagiarism_score, timestamp)
                        VALUES (?, ?, ?, ?)
                        ''', (st.session_state.current_student_id, uploaded_code.name, percentage, datetime.now()))
                        
                        conn.commit()
                        conn.close()
                    except Exception as e:
                        st.error(f"Error saving plagiarism check to database: {str(e)}")
### 📄 AI-Based Letter of Recommendation (LOR) Generator ###
if page == "📄 AI-Based LOR Generator":
    st.header("📄 AI-Powered Letter of Recommendation (LOR) Generator")

    # Select Student
    student_id = st.selectbox("Select Student ID", list(students_data.keys()))
    student = students_data[student_id]

    st.subheader(f"Student: {student['name']} | {student.get('department', 'B.Tech. Robotics & Artificial Intelligence')}, {student.get('year', 'Third Year')}")
    st.write(f"**CGPA:** {student.get('cgpa', 'Not Available')}")

    # Faculty Input: Purpose of LOR
    lor_purpose = st.radio("Purpose of LOR", ["Higher Studies", "Internship", "Job Application"])
    skills = st.text_area("Enter key skills (comma-separated)", "Machine Learning, Research, Leadership")
    achievements = st.text_area("Enter achievements (comma-separated)", "Won AI Hackathon, Published IEEE Paper")

    # Generate LOR
    if st.button("Generate LOR"):
        with st.spinner("Generating Letter of Recommendation..."):
            prompt = f"""
            Write a professional Letter of Recommendation (LOR) for a student. 
            The student details are:
            - Name: {student['name']}
            - Department: {student.get('department', 'B.Tech. Robotics & Artificial Intelligence, Third Year')}
            - CGPA: {student.get('cgpa', 'Not Available')}
            - Purpose: {lor_purpose if lor_purpose else "Not Provided"}
            - Skills: {skills if skills else "Not Provided"}
            - Achievements: {achievements if achievements else "Not Provided"}
  
            The LOR should be formal, structured, and highlight the student's strengths, academic achievements, and suitability for {lor_purpose}.
            """

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "system", "content": "You are an AI that writes professional letters of recommendation."},
                          {"role": "user", "content": prompt}],
                temperature=0.6,
                max_tokens=400
            )

            st.session_state.lor_text = response.choices[0].message.content.strip()
            st.success("✅ Letter of Recommendation Generated!")
            st.text_area("Generated LOR", st.session_state.lor_text, height=300)

    # Refinement System
    if "lor_text" in st.session_state:
        refine_prompt = st.text_area("🔄 Enter improvements to refine the LOR (optional)")

        if st.button("Refine LOR"):
            with st.spinner("Refining LOR..."):
                refine_request = f"Refine and improve the following LOR based on this feedback: {refine_prompt}\n\n{st.session_state.lor_text}"

                refined_response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "system", "content": "You refine and improve letters of recommendation based on feedback."},
                              {"role": "user", "content": refine_request}],
                    temperature=0.6,
                    max_tokens=400
                )

                st.session_state.lor_text = refined_response.choices[0].message.content.strip()
                st.success("✅ LOR Refined!")
                st.text_area("Refined LOR", st.session_state.lor_text, height=300)

    # Save and Download LOR as Word Document
    if "lor_text" in st.session_state:
        doc = docx.Document()
        doc.add_paragraph(st.session_state.lor_text)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(label="📥 Download LOR (Word)", data=buffer, file_name="LOR.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

